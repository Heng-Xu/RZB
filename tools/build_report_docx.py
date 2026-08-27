# -*- coding: utf-8 -*-
"""将研究报告 Markdown 底稿渲染为符合项目模板的 DOCX。

模板化用法：
    python tools/build_report_docx.py \
        --template 研究报告/中期研究报告初稿/03+XX项目-研究报告.docx \
        --input 研究报告/初稿/章节/研究报告第一至第三章初稿.md \
        --output 研究报告/初稿/研究报告第一至第三章初稿.docx \
        --title "项目名称" --subtitle "研究报告" \
        --organization "单位名称" --date "2026年8月"

兼容旧用法：
    python tools/build_report_docx.py <输入.md> <输出.docx>
    python tools/build_report_docx.py --onto <底稿.docx> <输入.md> <输出.docx>
"""

from __future__ import annotations

import argparse
import math
import re
import sys
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image


CJK = "仿宋_GB2312"
LATIN = "Times New Roman"
TITLE_CJK = "黑体"
TABLE_CJK = "仿宋_GB2312"
SONG = "宋体"

BODY_SIZE = 14
BODY_LINE_TWIPS = 560  # 固定 28 pt，与项目模板一致
TABLE_SIZE = 12
CAPTION_SIZE = 12


def _set_rfonts(rpr, cjk: str, latin: str) -> None:
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), cjk)


def set_run_fonts(
    run,
    size_pt: float,
    *,
    bold: bool = False,
    cjk: str = CJK,
    latin: str = LATIN,
) -> None:
    """为中文、英文和数字分别设置字体，避免中文回退为宋体。"""

    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    _set_rfonts(rpr, cjk, latin)


def _set_exact_line_spacing(paragraph, line_twips: int = BODY_LINE_TWIPS) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    spacing.set(qn("w:line"), str(line_twips))
    spacing.set(qn("w:lineRule"), "exact")


def set_first_line_indent(paragraph, chars: int = 2) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    ind.set(qn("w:firstLineChars"), str(chars * 100))
    # 模板同时保留 firstLine=200；Word/WPS 以字符缩进为主。
    ind.set(qn("w:firstLine"), str(chars * 100))
    ind.attrib.pop(qn("w:hanging"), None)
    ind.attrib.pop(qn("w:hangingChars"), None)


def _set_no_indent(paragraph) -> None:
    """显式覆盖 Normal 样式的首行缩进，供封面、图题等段落使用。"""

    ppr = paragraph._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    ind.set(qn("w:firstLine"), "0")
    ind.set(qn("w:firstLineChars"), "0")
    ind.attrib.pop(qn("w:hanging"), None)
    ind.attrib.pop(qn("w:hangingChars"), None)


def _set_hanging_indent(paragraph, chars: int = 2) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    ind.attrib.pop(qn("w:firstLine"), None)
    ind.attrib.pop(qn("w:firstLineChars"), None)
    ind.set(qn("w:leftChars"), str(chars * 100))
    ind.set(qn("w:hangingChars"), str(chars * 100))
    ind.set(qn("w:left"), str(chars * BODY_SIZE * 20))
    ind.set(qn("w:hanging"), str(chars * BODY_SIZE * 20))


def _style_font(style, size: float, *, bold: bool, cjk: str = CJK) -> None:
    style.font.name = LATIN
    style.font.size = Pt(size)
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    _set_rfonts(rpr, cjk, LATIN)


def _style_spacing(style, *, line_twips: int, before: int = 0, after: int = 0) -> None:
    ppr = style.element.get_or_add_pPr()
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    spacing.set(qn("w:line"), str(line_twips))
    spacing.set(qn("w:lineRule"), "exact")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))


def setup_page_and_styles(doc: Document) -> None:
    """按项目模板统一页面、正文、标题、图题和表格样式。"""

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    _style_font(normal, BODY_SIZE, bold=False)
    _style_spacing(normal, line_twips=BODY_LINE_TWIPS)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal_ppr = normal.element.get_or_add_pPr()
    normal_ind = normal_ppr.find(qn("w:ind"))
    if normal_ind is None:
        normal_ind = OxmlElement("w:ind")
        normal_ppr.append(normal_ind)
    normal_ind.set(qn("w:firstLine"), "200")
    normal_ind.set(qn("w:firstLineChars"), "200")

    heading_specs = {
        "Heading 1": (18, 560, 120, 120),
        "Heading 2": (14, 560, 120, 120),
        "Heading 3": (14, 560, 50, 50),
    }
    for name, (size, line, before, after) in heading_specs.items():
        style = doc.styles[name]
        _style_font(style, size, bold=True)
        _style_spacing(style, line_twips=line, before=before, after=after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.first_line_indent = None
        # 标题文字已经含章、节编号，移除模板自动编号以免重复显示。
        style_ppr = style.element.get_or_add_pPr()
        numpr = style_ppr.find(qn("w:numPr"))
        if numpr is not None:
            style_ppr.remove(numpr)

    if "Caption" in [style.name for style in doc.styles]:
        caption = doc.styles["Caption"]
        _style_font(caption, CAPTION_SIZE, bold=False)
        caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.first_line_indent = None

    if "表" in [style.name for style in doc.styles]:
        table_style = doc.styles["表"]
        _style_font(table_style, TABLE_SIZE, bold=False, cjk=TABLE_CJK)

    for level in (1, 2):
        name = f"toc {level}"
        if name not in [style.name for style in doc.styles]:
            continue
        toc_style = doc.styles[name]
        _style_font(toc_style, 14, bold=level == 1)
        _style_spacing(toc_style, line_twips=BODY_LINE_TWIPS)


def _clear_template_body(doc: Document) -> None:
    """删除模板示例正文，保留样式、页面设置及最后一个节属性。"""

    body = doc._element.body
    final_sectpr = body.sectPr
    for child in list(body):
        if child is not final_sectpr:
            body.remove(child)


def _split_bold(text: str, base_bold: bool = False):
    parts = text.split("**")
    segments = []
    for index, part in enumerate(parts):
        if part:
            segments.append((part, base_bold or index % 2 == 1))
    return segments or [(text, base_bold)]


def _clean_inline(text: str) -> str:
    """清理 Pandoc/GFM 为普通字符增加的转义，不改变正文内容。"""

    for escaped, plain in (
        (r"\[", "["),
        (r"\]", "]"),
        (r"\<", "<"),
        (r"\>", ">"),
        (r"\|", "|"),
        (r"\_", "_"),
    ):
        text = text.replace(escaped, plain)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def _display_units(text: str) -> float:
    """估算混排文字所占的中文全角字符数，用于目录缓存页码初值。"""

    units = 0.0
    for char in _clean_inline(text).replace("**", ""):
        if char.isspace():
            units += 0.5
        elif unicodedata.east_asian_width(char) in {"W", "F", "A"}:
            units += 1.0
        else:
            units += 0.55
    return units


def _extract_toc_entries(md_path: Path) -> list[dict[str, object]]:
    entries: list[dict[str, object]] = []
    for line in md_path.read_text(encoding="utf-8-sig").splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            level = 1
            text = _clean_inline(stripped[2:])
        elif stripped.startswith("## "):
            level = 2
            text = _clean_inline(stripped[3:])
        else:
            continue
        entries.append(
            {
                "level": level,
                "text": text,
                "bookmark": f"_TocReport{len(entries) + 1:04d}",
                "bookmark_id": len(entries) + 1,
            }
        )
    return entries


def _estimate_toc_pages(md_path: Path, entries: list[dict[str, object]]) -> list[int]:
    """按范文 14 pt/28 pt 版式估算目录缓存页码；PAGEREF 打开后可更新。"""

    lines = md_path.read_text(encoding="utf-8-sig").splitlines()
    page_capacity = 675.0
    text_capacity = 29.5
    page = 1
    used = 0.0
    first_chapter = True
    entry_pages: list[int] = []

    def advance(height: float, *, keep_together: bool = False) -> None:
        nonlocal page, used
        if keep_together and used > 0 and height > page_capacity - used:
            page += 1
            used = 0.0
        remaining_height = height
        while remaining_height > page_capacity - used + 0.01:
            remaining_height -= max(page_capacity - used, 0.0)
            page += 1
            used = 0.0
        used += remaining_height

    index = 0
    entry_index = 0
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            index += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while index < len(lines):
                candidate = lines[index].strip()
                if not (candidate.startswith("|") and candidate.endswith("|")):
                    break
                table_lines.append(candidate)
                index += 1
            rows = _parse_table_rows(table_lines)
            columns = max((len(row) for row in rows), default=1)
            cell_capacity = max(5.0, (415.0 / columns - 8.0) / TABLE_SIZE)
            table_height = 0.0
            for row in rows:
                wrapped = max(
                    (max(1, math.ceil(_display_units(cell) / cell_capacity)) for cell in row),
                    default=1,
                )
                table_height += wrapped * 18.0 + 4.0
            advance(table_height)
            continue

        image_match = re.fullmatch(r"!\[(.+?)\]\((.+?)\)", stripped)
        if image_match:
            image_path = (md_path.parent / _clean_inline(image_match.group(2)).strip("<>")).resolve()
            with Image.open(image_path) as image:
                picture_height = 391.2 * image.height / image.width
            advance(picture_height + 28.0, keep_together=True)
            index += 1
            continue

        if stripped.startswith("# "):
            if first_chapter:
                first_chapter = False
            else:
                page += 1
                used = 0.0
            entry_pages.append(page)
            entry_index += 1
            advance(40.0, keep_together=True)
        elif stripped.startswith("## "):
            if page_capacity - used < 68.0:
                page += 1
                used = 0.0
            entry_pages.append(page)
            entry_index += 1
            advance(40.0, keep_together=True)
        elif stripped.startswith(("### ", "#### ")):
            advance(33.0, keep_together=True)
        elif re.fullmatch(r"\*\*表\s+.+\*\*", stripped):
            advance(24.0, keep_together=True)
        else:
            line_count = max(1, math.ceil((_display_units(stripped) + 2.0) / text_capacity))
            advance(line_count * 28.0)
        index += 1

    if entry_index != len(entries):
        raise ValueError("目录条目与正文章、节标题数量不一致")
    return entry_pages


def _add_runs(paragraph, text: str, size: float, *, bold: bool = False, cjk: str = CJK):
    for segment, segment_bold in _split_bold(_clean_inline(text), bold):
        run = paragraph.add_run(segment)
        set_run_fonts(run, size, bold=segment_bold, cjk=cjk)


def add_paragraph(
    doc: Document,
    text: str,
    *,
    size: float = BODY_SIZE,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    indent: bool = True,
    hanging: bool = False,
    style: str = "Normal",
    cjk: str = CJK,
    line_twips: int = BODY_LINE_TWIPS,
    space_before: float = 0,
    space_after: float = 0,
):
    paragraph = doc.add_paragraph(style=style)
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    _set_exact_line_spacing(paragraph, line_twips)
    if indent:
        set_first_line_indent(paragraph)
    else:
        _set_no_indent(paragraph)
    if hanging:
        _set_hanging_indent(paragraph)
    _add_runs(paragraph, text, size, bold=bold, cjk=cjk)
    return paragraph


def _add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    insert_at = 1 if paragraph._p.pPr is not None else 0
    paragraph._p.insert(insert_at, start)
    paragraph._p.append(end)


def _add_heading(
    doc: Document,
    text: str,
    level: int,
    *,
    page_break_before: bool = False,
    bookmark_name: str | None = None,
    bookmark_id: int | None = None,
):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.page_break_before = page_break_before
    size = 18 if level == 1 else 14
    _add_runs(paragraph, text, size, bold=True)
    if bookmark_name is not None and bookmark_id is not None:
        _add_bookmark(paragraph, bookmark_name, bookmark_id)
    return paragraph


def _add_cover(
    doc: Document,
    *,
    title: str,
    subtitle: str,
    organization: str,
    report_date: str,
) -> None:
    for _ in range(4):
        doc.add_paragraph()
    add_paragraph(
        doc,
        title,
        size=22,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        indent=False,
        cjk=TITLE_CJK,
        line_twips=640,
        space_after=18,
    )
    add_paragraph(
        doc,
        subtitle,
        size=22,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        indent=False,
        cjk=TITLE_CJK,
        line_twips=640,
    )
    add_paragraph(
        doc,
        organization,
        size=16,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        indent=False,
        cjk=SONG,
        line_twips=580,
        space_before=250,
    )
    add_paragraph(
        doc,
        report_date,
        size=16,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        indent=False,
        cjk=SONG,
        line_twips=580,
    )


def _add_field(
    paragraph,
    instruction: str,
    *,
    result_text: str | None = None,
    dirty: bool = False,
    size: float = 14,
) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    if dirty:
        begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate])
    if result_text is not None:
        result = paragraph.add_run(result_text)
        set_run_fonts(result, size, bold=False)
    end_run = paragraph.add_run()
    end_run._r.append(end)


def _set_right_dot_tab(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    tabs = ppr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        ppr.append(tabs)
    for existing in list(tabs):
        tabs.remove(existing)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), "8200")
    tabs.append(tab)


def _add_toc(doc: Document, entries: list[dict[str, object]], page_hints: list[int]) -> None:
    add_paragraph(
        doc,
        "目  录",
        size=18,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        indent=False,
        cjk=TITLE_CJK,
        line_twips=560,
        space_after=18,
    )
    available_styles = [style.name for style in doc.styles]
    for entry, page_hint in zip(entries, page_hints):
        level = int(entry["level"])
        style = f"toc {level}" if f"toc {level}" in available_styles else "Normal"
        paragraph = doc.add_paragraph(style=style)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_no_indent(paragraph)
        paragraph.paragraph_format.left_indent = Cm(0.74 * (level - 1))
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        _set_exact_line_spacing(paragraph, BODY_LINE_TWIPS)
        _set_right_dot_tab(paragraph)
        _add_runs(paragraph, str(entry["text"]), 14, bold=level == 1)
        tab_run = paragraph.add_run()
        set_run_fonts(tab_run, 14, bold=False)
        tab_run.add_tab()
        _add_field(
            paragraph,
            f" PAGEREF {entry['bookmark']} \\h ",
            result_text=str(page_hint),
            dirty=True,
            size=14,
        )


def _start_body_section(doc: Document):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.different_first_page_header_footer = False
    pg_num = section._sectPr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        section._sectPr.append(pg_num)
    pg_num.set(qn("w:start"), "1")
    return section


def _enable_field_updates(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def _add_page_number(doc: Document, *, body_section_index: int | None = None) -> None:
    for index, section in enumerate(doc.sections):
        section.footer.is_linked_to_previous = False
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.clear()
        _set_no_indent(paragraph)
        if body_section_index is None:
            section.different_first_page_header_footer = True
            _add_field(paragraph, " PAGE ", result_text="1", dirty=True, size=9)
        elif index == body_section_index:
            section.different_first_page_header_footer = False
            _add_field(paragraph, " PAGE ", result_text="1", dirty=True, size=9)


def _is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def _parse_table_rows(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        if _is_table_separator(line):
            continue
        cells = [_clean_inline(cell) for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def _set_cell_text(cell, text: str, *, header: bool, first_column: bool) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.style = "表" if "表" in [style.name for style in paragraph.part.document.styles] else "Normal"
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if header or not first_column else WD_ALIGN_PARAGRAPH.LEFT
    )
    _set_no_indent(paragraph)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    _set_exact_line_spacing(paragraph, 360)
    _add_runs(paragraph, text, TABLE_SIZE, bold=header, cjk=TABLE_CJK)


def _set_table_borders(table) -> None:
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def _add_table(doc: Document, raw_lines: list[str]) -> None:
    rows = _parse_table_rows(raw_lines)
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.autofit = True
    _set_table_borders(table)
    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            value = row[column_index] if column_index < len(row) else ""
            _set_cell_text(
                table.cell(row_index, column_index),
                value,
                header=row_index == 0,
                first_column=column_index == 0,
            )
        row_pr = table.rows[row_index]._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        row_pr.append(cant_split)
        if row_index == 0:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            row_pr.append(repeat)


def _add_image(doc: Document, md_path: Path, alt_text: str, image_ref: str) -> None:
    image_ref = _clean_inline(image_ref).strip("<>")
    image_path = (md_path.parent / image_ref).resolve()
    if not image_path.exists():
        raise FileNotFoundError(f"Markdown 引用的图片不存在：{image_path}")
    paragraph = doc.add_paragraph(style="图片" if "图片" in [s.name for s in doc.styles] else "Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_no_indent(paragraph)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(13.8))
    caption = add_paragraph(
        doc,
        _clean_inline(alt_text),
        size=CAPTION_SIZE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        indent=False,
        style="Caption" if "Caption" in [s.name for s in doc.styles] else "Normal",
        line_twips=400,
        space_after=6,
    )
    caption.paragraph_format.keep_with_next = False


def render_md_into(
    doc: Document,
    md_path,
    *,
    toc_entries: list[dict[str, object]] | None = None,
) -> Document:
    """把 GFM Markdown 正文渲染为标题、段落、表格、图片和参考文献。"""

    md_path = Path(md_path)
    lines = md_path.read_text(encoding="utf-8-sig").splitlines()
    in_references = False
    seen_level_one = False
    toc_entry_index = 0
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        if stripped.startswith("|" ) and stripped.endswith("|"):
            table_lines = []
            while index < len(lines):
                candidate = lines[index].strip()
                if not (candidate.startswith("|") and candidate.endswith("|")):
                    break
                table_lines.append(candidate)
                index += 1
            _add_table(doc, table_lines)
            continue

        image_match = re.fullmatch(r"!\[(.+?)\]\((.+?)\)", stripped)
        if image_match:
            _add_image(doc, md_path, image_match.group(1), image_match.group(2))
            index += 1
            continue

        if stripped.startswith("# "):
            heading_text = _clean_inline(stripped[2:])
            in_references = heading_text == "参考文献"
            entry = None
            if toc_entries is not None:
                entry = toc_entries[toc_entry_index]
                if int(entry["level"]) != 1 or str(entry["text"]) != heading_text:
                    raise ValueError("目录条目与一级标题不一致")
                toc_entry_index += 1
            _add_heading(
                doc,
                heading_text,
                1,
                page_break_before=seen_level_one,
                bookmark_name=str(entry["bookmark"]) if entry else None,
                bookmark_id=int(entry["bookmark_id"]) if entry else None,
            )
            seen_level_one = True
        elif stripped.startswith("## "):
            heading_text = _clean_inline(stripped[3:])
            entry = None
            if toc_entries is not None:
                entry = toc_entries[toc_entry_index]
                if int(entry["level"]) != 2 or str(entry["text"]) != heading_text:
                    raise ValueError("目录条目与二级标题不一致")
                toc_entry_index += 1
            _add_heading(
                doc,
                heading_text,
                2,
                bookmark_name=str(entry["bookmark"]) if entry else None,
                bookmark_id=int(entry["bookmark_id"]) if entry else None,
            )
        elif stripped.startswith("### "):
            _add_heading(doc, _clean_inline(stripped[4:]), 3)
        elif stripped.startswith("#### "):
            _add_heading(doc, _clean_inline(stripped[5:]), 3)
        elif re.fullmatch(r"\*\*表\s+.+\*\*", stripped):
            caption_text = _clean_inline(stripped[2:-2])
            caption = add_paragraph(
                doc,
                caption_text,
                size=CAPTION_SIZE,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                indent=False,
                style="Caption" if "Caption" in [s.name for s in doc.styles] else "Normal",
                line_twips=400,
                space_before=6,
                space_after=3,
            )
            caption.paragraph_format.keep_with_next = True
        elif in_references and re.match(r"^\\?\[\d+\]", stripped):
            add_paragraph(
                doc,
                _clean_inline(stripped),
                size=BODY_SIZE,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                indent=False,
                hanging=True,
                line_twips=BODY_LINE_TWIPS,
            )
        elif re.match(r"^(R\(|K_[A-Za-z]|CRF\()", stripped):
            add_paragraph(
                doc,
                _clean_inline(stripped),
                size=BODY_SIZE,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                indent=False,
            )
        elif stripped.startswith(("- ", "* ")):
            add_paragraph(doc, "• " + stripped[2:], size=BODY_SIZE, indent=False)
        elif stripped.startswith("> "):
            add_paragraph(doc, stripped[2:], size=BODY_SIZE, indent=False)
        elif set(stripped) <= {"-"} and len(stripped) >= 3:
            pass
        else:
            add_paragraph(doc, stripped, size=BODY_SIZE, indent=True)
        index += 1
    if toc_entries is not None and toc_entry_index != len(toc_entries):
        raise ValueError("目录条目未全部写入正文书签")
    return doc


def build_report(
    *,
    template_path,
    md_path,
    out_path,
    title: str,
    subtitle: str,
    organization: str,
    report_date: str,
    include_toc: bool = True,
):
    """以项目模板样式生成含封面、目录、正文、图表和页码的报告。"""

    template_path = Path(template_path)
    md_path = Path(md_path)
    out_path = Path(out_path)
    toc_entries = _extract_toc_entries(md_path) if include_toc else []
    page_hints = _estimate_toc_pages(md_path, toc_entries) if include_toc else []
    doc = Document(template_path)
    _clear_template_body(doc)
    setup_page_and_styles(doc)
    _add_cover(
        doc,
        title=title,
        subtitle=subtitle,
        organization=organization,
        report_date=report_date,
    )
    doc.add_page_break()
    if include_toc:
        _add_toc(doc, toc_entries, page_hints)
        _start_body_section(doc)
    render_md_into(doc, md_path, toc_entries=toc_entries if include_toc else None)
    _add_page_number(doc, body_section_index=1 if include_toc else None)
    _enable_field_updates(doc)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def build_standalone(md_path, out_path):
    """兼容旧调用：不用模板，直接生成正文文档。"""

    doc = Document()
    setup_page_and_styles(doc)
    render_md_into(doc, md_path)
    _add_page_number(doc)
    _enable_field_updates(doc)
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def build_onto(base_docx, md_path, out_path):
    """兼容旧调用：在既有文档后分页追加 Markdown 正文。"""

    doc = Document(base_docx)
    setup_page_and_styles(doc)
    doc.add_page_break()
    render_md_into(doc, md_path)
    _add_page_number(doc)
    _enable_field_updates(doc)
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def _parse_args(argv: list[str]):
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--template", type=Path)
    parser.add_argument("--input", dest="input_path", type=Path)
    parser.add_argument("--output", type=Path)
    parser.add_argument("--title")
    parser.add_argument("--subtitle", default="研究报告")
    parser.add_argument("--organization", default="")
    parser.add_argument("--date", dest="report_date", default="")
    parser.add_argument("--no-toc", action="store_true", help="不生成实体目录；默认生成章、节两级目录")
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    argv = list(sys.argv[1:] if argv is None else argv)
    if argv and argv[0] == "--onto":
        if len(argv) != 4:
            raise SystemExit("用法：build_report_docx.py --onto <base.docx> <input.md> <out.docx>")
        output = build_onto(argv[1], argv[2], argv[3])
    elif argv and not argv[0].startswith("-"):
        if len(argv) != 2:
            raise SystemExit("用法：build_report_docx.py <input.md> <out.docx>")
        output = build_standalone(argv[0], argv[1])
    else:
        args = _parse_args(argv)
        required = {
            "--template": args.template,
            "--input": args.input_path,
            "--output": args.output,
            "--title": args.title,
        }
        missing = [name for name, value in required.items() if not value]
        if missing:
            raise SystemExit("缺少参数：" + "、".join(missing))
        output = build_report(
            template_path=args.template,
            md_path=args.input_path,
            out_path=args.output,
            title=args.title,
            subtitle=args.subtitle,
            organization=args.organization,
            report_date=args.report_date,
            include_toc=not args.no_toc,
        )
    print("WROTE:", output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
