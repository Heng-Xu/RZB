# -*- coding: utf-8 -*-
"""
将含 Markdown 表格的底稿渲染为 DOCX（表格渲染为真正的 Word 表格）。
排版：A4、页边距 左右2.5cm/上下2.54cm、中文 仿宋_GB2312、数字/英文 Times New Roman。
层级：# 标题(二号粗居中) / ## 章(三号粗左) / ### 节(四号粗左) /
      表格(自动识别 | a | b | 行 + |---| 分隔行，表头加粗) /
      - 列表 / > 引用 / 正文(小四) / 行内 **粗体**。
用法：python tools/build_md_table_docx.py <输入md> <输出docx>
"""
import sys, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CJK = "仿宋_GB2312"
LATIN = "Times New Roman"


def set_run_fonts(run, size_pt, bold=False, cjk=CJK, latin=LATIN, color=None):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), latin)
    rfonts.set(qn('w:hAnsi'), latin)
    rfonts.set(qn('w:eastAsia'), cjk)
    if color is not None:
        run.font.color.rgb = color


def split_bold(text):
    """按 **粗体** 切分为 (片段, 是否粗体)。"""
    parts, last, bold = [], 0, False
    for m in re.finditer(r'\*\*(.+?)\*\*', text):
        if m.start() > last:
            parts.append((text[last:m.start()], False))
        parts.append((m.group(1), True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False))
    return parts or [(text, False)]


def add_para(doc, text, size, bold=False, align=None, indent=False,
             line_spacing=1.5, space_before=0, space_after=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent:
        pf.first_line_indent = Pt(size * 2)
    for frag, b in split_bold(text):
        if not frag:
            continue
        run = p.add_run(frag)
        set_run_fonts(run, size, bold=bold or b)
    if not p.runs:
        set_run_fonts(p.add_run(''), size, bold=bold)
    return p


def set_cell(cell, text, size=9, bold=False, header=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for frag, b in split_bold(text.strip()):
        if not frag:
            continue
        run = p.add_run(frag)
        set_run_fonts(run, size, bold=bold or b)
    if not p.runs:
        set_run_fonts(p.add_run(''), size, bold=bold)
    if header:
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), 'D9E2F3')
        cell._tc.get_or_add_tcPr().append(shd)


def add_table(doc, rows):
    """rows: list of list[str]，第一行为表头。"""
    ncol = max(len(r) for r in rows)
    rows = [r + [''] * (ncol - len(r)) for r in rows]
    tbl = doc.add_table(rows=len(rows), cols=ncol)
    tbl.style = 'Table Grid'
    tbl.autofit = True
    for i, row in enumerate(rows):
        for j, txt in enumerate(row):
            set_cell(tbl.rows[i].cells[j], txt, size=9,
                     bold=(i == 0), header=(i == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def parse_table_block(lines, idx):
    """从 idx 起收集连续表格行，返回 (rows, next_idx) 或 (None, idx)。"""
    block = []
    j = idx
    while j < len(lines) and lines[j].strip().startswith('|'):
        block.append(lines[j].strip())
        j += 1
    if len(block) < 2:
        return None, idx
    rows = []
    for k, ln in enumerate(block):
        cells = [c.strip() for c in ln.strip('|').split('|')]
        if k == 1 and all(set(c) <= set('-: ') for c in cells if c):
            continue  # 分隔行
        rows.append(cells)
    return rows, j


def setup_page(doc):
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Cm(29.7), Cm(21.0)
    sec.left_margin = sec.right_margin = Cm(1.8)
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    normal = doc.styles['Normal']
    normal.font.name = LATIN
    normal.font.size = Pt(12)
    normal._element.get_or_add_rPr().append(
        OxmlElement('w:rFonts'))
    rfonts = normal._element.rPr.find(qn('w:rFonts'))
    rfonts.set(qn('w:ascii'), LATIN)
    rfonts.set(qn('w:hAnsi'), LATIN)
    rfonts.set(qn('w:eastAsia'), CJK)


def build(md_path, out_path):
    doc = Document()
    setup_page(doc)
    with open(md_path, encoding='utf-8') as f:
        lines = f.read().split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        s = line.strip()
        if s.startswith('|'):
            rows, nxt = parse_table_block(lines, i)
            if rows:
                add_table(doc, rows)
                i = nxt
                continue
        if not s:
            i += 1
            continue
        if s.startswith('# '):
            add_para(doc, s[2:], 18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                     space_before=4, space_after=12, line_spacing=1.2)
        elif s.startswith('## '):
            add_para(doc, s[3:], 15, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
                     space_before=12, space_after=6, line_spacing=1.2)
        elif s.startswith('### '):
            add_para(doc, s[4:], 13, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
                     space_before=8, space_after=4, line_spacing=1.2)
        elif set(s) <= set('-') and len(s) >= 3:
            i += 1
            continue  # 分割线
        elif s == '>':
            i += 1
            continue
        elif s.startswith('> '):
            add_para(doc, s[2:], 11, indent=False, line_spacing=1.25,
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=2)
        elif re.match(r'^\d+\.\s', s):
            add_para(doc, '　' + s, 12, indent=False, line_spacing=1.4,
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=0)
        elif s.startswith('- ') or s.startswith('* '):
            add_para(doc, '　• ' + s[2:], 12, indent=False, line_spacing=1.4,
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=0)
        else:
            add_para(doc, s, 12, indent=True, line_spacing=1.5,
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=0)
        i += 1
    doc.save(out_path)
    return out_path


if __name__ == '__main__':
    out = build(sys.argv[1], sys.argv[2])
    print('OK', out)
