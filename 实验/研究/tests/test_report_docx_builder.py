# -*- coding: utf-8 -*-
"""研究报告 Word 生成脚本的模板、结构和版式回归测试。"""

from __future__ import annotations

import importlib.util
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image


PROJECT_ROOT = Path(__file__).resolve().parents[3]
SCRIPT_PATH = PROJECT_ROOT / "tools" / "build_report_docx.py"
TEMPLATE_PATH = (
    PROJECT_ROOT
    / "研究报告"
    / "中期研究报告初稿"
    / "03+XX项目-研究报告.docx"
)


def _load_builder():
    spec = importlib.util.spec_from_file_location("build_report_docx", SCRIPT_PATH)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def _write_fixture(tmp_path: Path) -> Path:
    # 测试只核验图片是否真正进入 docx，不依赖报告正式图件。
    Image.new("RGB", (16, 16), "white").save(tmp_path / "figure.png")
    md = tmp_path / "fixture.md"
    md.write_text(
        """# 第一章 绪论

## 1.1 研究背景与意义

### 1.1.1 测试小节

这是正文段落，包含 **重点内容** 和 Times New Roman 123。

**表 1-1 测试数据表**

| 指标 | 数值 |
|---|---:|
| 容载比 | 2.24 |

![图 1-1 测试图件](figure.png)

# 参考文献

[1] 测试文献[J]. 2026.
""",
        encoding="utf-8",
    )
    return md


def _build_fixture(tmp_path: Path) -> Path:
    builder = _load_builder()
    md = _write_fixture(tmp_path)
    output = tmp_path / "report.docx"
    builder.build_report(
        template_path=TEMPLATE_PATH,
        md_path=md,
        out_path=output,
        title="徐州地区分布式新能源高渗透率地区110kV电网容载比弹性指标优化研究",
        subtitle="研究报告（第一至第三章初稿）",
        organization="国网江苏省电力有限公司徐州供电分公司",
        report_date="2026年8月",
        include_toc=True,
    )
    return output


def test_builder_renders_cover_headings_table_image_and_reference(tmp_path):
    output = _build_fixture(tmp_path)
    assert output.exists() and output.stat().st_size > 20_000

    doc = Document(output)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert texts[0] == "徐州地区分布式新能源高渗透率地区110kV电网容载比弹性指标优化研究"
    assert "研究报告（第一至第三章初稿）" in texts
    assert "目  录" in texts
    toc_entries = [p for p in doc.paragraphs if p.style.name.lower().startswith("toc")]
    assert any(p.text.startswith("第一章 绪论\t") for p in toc_entries)
    assert any(p.text.startswith("1.1 研究背景与意义\t") for p in toc_entries)
    assert "第一章 绪论" in texts
    assert "1.1 研究背景与意义" in texts
    assert "1.1.1 测试小节" in texts
    assert "表 1-1 测试数据表" in texts
    assert "图 1-1 测试图件" in texts
    assert "[1] 测试文献[J]. 2026." in texts

    assert len(doc.tables) == 1
    assert doc.tables[0].cell(0, 0).text == "指标"
    assert doc.tables[0].cell(1, 1).text == "2.24"
    assert "w:drawing" in doc._element.xml


def test_builder_matches_template_page_and_typography(tmp_path):
    output = _build_fixture(tmp_path)
    doc = Document(output)
    section = doc.sections[0]
    assert round(section.page_width.cm, 2) == 21.00
    assert round(section.page_height.cm, 2) == 29.70
    assert round(section.left_margin.cm, 2) == 3.17
    assert round(section.right_margin.cm, 2) == 3.17
    assert round(section.top_margin.cm, 2) == 2.54
    assert round(section.bottom_margin.cm, 2) == 2.54

    cover = next(p for p in doc.paragraphs if p.text.startswith("徐州地区"))
    assert cover.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert cover.runs[0].font.size.pt == 22
    assert cover.runs[0].bold is True
    assert cover.runs[0]._r.rPr.rFonts.get(qn("w:eastAsia")) == "黑体"
    assert cover._p.pPr.ind.get(qn("w:firstLineChars")) == "0"

    chapter = next(p for p in doc.paragraphs if p.text == "第一章 绪论")
    section_heading = next(p for p in doc.paragraphs if p.text == "1.1 研究背景与意义")
    subsection = next(p for p in doc.paragraphs if p.text == "1.1.1 测试小节")
    assert chapter.style.name == "Heading 1"
    assert chapter.runs[0].font.size.pt == 18
    assert doc.styles["Heading 1"].element.pPr.numPr is None
    assert section_heading.style.name == "Heading 2"
    assert section_heading.runs[0].font.size.pt == 14
    assert subsection.style.name == "Heading 3"
    assert subsection.runs[0].font.size.pt == 14

    body = next(p for p in doc.paragraphs if p.text.startswith("这是正文段落"))
    assert body.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert body.runs[0].font.size.pt == 14
    assert body.runs[0]._r.rPr.rFonts.get(qn("w:eastAsia")) == "仿宋_GB2312"
    assert body.runs[0]._r.rPr.rFonts.get(qn("w:ascii")) == "Times New Roman"
    assert body._p.pPr.ind.get(qn("w:firstLineChars")) == "200"
    assert body._p.pPr.spacing.get(qn("w:line")) == "560"
    assert body._p.pPr.spacing.get(qn("w:lineRule")) == "exact"

    table_run = doc.tables[0].cell(1, 0).paragraphs[0].runs[0]
    assert table_run.font.size.pt == 12
    assert table_run._r.rPr.rFonts.get(qn("w:eastAsia")) == "仿宋_GB2312"


def test_builder_inserts_visible_pageref_toc_and_restarts_body_page_number(tmp_path):
    output = _build_fixture(tmp_path)
    with ZipFile(output) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        settings_xml = archive.read("word/settings.xml").decode("utf-8")
        footer_xml = "".join(
            archive.read(name).decode("utf-8")
            for name in archive.namelist()
            if name.startswith("word/footer") and name.endswith(".xml")
        )

    assert "PAGEREF" in document_xml
    assert "第一章 绪论" in document_xml
    assert "updateFields" in settings_xml
    assert "PAGE" in footer_xml

    doc = Document(output)
    assert len(doc.sections) == 2
    page_numbering = doc.sections[1]._sectPr.find(qn("w:pgNumType"))
    assert page_numbering is not None
    assert page_numbering.get(qn("w:start")) == "1"
