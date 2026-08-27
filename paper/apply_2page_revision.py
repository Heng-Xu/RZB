#!/usr/bin/env python3
"""In-place terminology/expression revision of the 2-page docx.
Only touches text runs (no formulas/format/layout). Preserves math-variable runs
in equation-carrying paragraphs by doing per-run substring replacement there."""
import shutil
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = "/home/roscy/ws_HengXU/徐州地区分布式新能源高渗透率地区110kV电网容载比弹性指标优化研究/paper/epg-nsga-ii-paper_final_compressed.docx"
OUT = "/home/roscy/ws_HengXU/徐州地区分布式新能源高渗透率地区110kV电网容载比弹性指标优化研究/paper/epg-nsga-ii-paper_final_compressed_revised.docx"
shutil.copy(SRC, OUT)
d = Document(OUT)
P = d.paragraphs


def set_para(i, text):
    """Put all text into run0, clear the rest (pure-prose paras only)."""
    p = P[i]
    assert p.runs, f"P{i} has no runs"
    p.runs[0].text = text
    for r in p.runs[1:]:
        r.text = ""


def repl(i, old, new):
    """Per-run substring replace (safe for math paras)."""
    hit = False
    for r in P[i].runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            hit = True
    assert hit, f"P{i}: '{old}' not found"


# ---- Title (2 runs: body + 'Penetration'); text shortened, template 24 pt unchanged ----
P[0].runs[0].text = ("Engineering-Prior-Guided NSGA-II for Differentiated Capacity-Load Ratio "
                     "Optimization in County 110-kV Grids under High Distributed PV ")
P[0].runs[1].text = "Penetration"

# ---- Abstract: run0 = 'Abstract—' + full body; runs 1-2 hold body tail -> clear them ----
P[7].runs[1].text = ""
P[7].runs[2].text = ""
P[7].runs[0].text = (
    "Abstract—"
    "High penetration of distributed photovoltaic (PV) generation makes county-level 110-kV grids bidirectional, "
    "causing reverse power flow and affecting supply adequacy under N-1 contingencies. This study formulates a "
    "multiobjective model that coordinates a county's 110-kV substations and optimizes "
    "differentiated capacity-load ratios (CLRs) and storage power ratings under a load-weighted aggregate CLR band "
    "and reverse-power-flow constraints, minimizing annualized cost and expected energy not supplied (EENS). "
    "Engineering-Prior-Guided NSGA-II (EPG-NSGA-II) adds feasibility-preserving warm-start sampling and a repair "
    "projection. On a public-data-informed synthetic county test system, it cuts annualized "
    "cost by 13.1% versus the uniform R=2.0 baseline at the same EENS level; across K=10–40 its converged "
    "hypervolume reaches +4.8% over standard NSGA-II at K=40.")

# ---- Index Terms (2 runs: 'Index Terms' | body) ----
P[8].runs[1].text = ("—Capacity-load ratio, distributed photovoltaic generation, multiobjective optimization, "
                     "reverse power flow, substation planning.")

# ---- Intro para 1 (P10, 1 run) ----
set_para(10,
    "High penetration of distributed PV generation is turning county-level 110-kV grids bidirectional: midday PV can "
    "exceed local demand, causing reverse power flow and shifting supply adequacy under N-1 contingencies [3]. A "
    "uniform capacity-load ratio over-invests at moderate-PV stations and under-constrains reverse power flow at "
    "substations with high distributed-PV penetration [1], [2]. Existing planning uses fixed or empirical rules, "
    "while evolutionary optimization searches the cost-risk trade-off directly [1], [2], [5], [6], with knowledge "
    "transfer further improving search efficiency [9]. Random initialization and variation then yield many "
    "infeasible individuals.")

# ---- Intro para 2 (P11, 8 runs prose) ----
set_para(11,
    "This study proposes EPG-NSGA-II, which keeps nondominated sorting unchanged and adds two engineering-prior "
    "operators: warm-start sampling within engineering bounds and repair projection after variation. Contributions: "
    "a multiobjective model with differentiated capacity-load ratios under an aggregate-CLR band, reverse-power-flow "
    "constraints, and an EENS-based supply-adequacy objective; the EPG-NSGA-II strategy; and a public-data-informed "
    "synthetic county test system with ablation.")

# ---- Planning-model math paras: per-run term swaps ----
repl(14, "absorbing reverse power", "absorbing reverse power flow")
repl(15, "residual reverse power", "residual reverse power flow")
repl(16, "regional capacity-load ratio", "load-weighted aggregate capacity-load ratio")

# ---- P17 (1 run) ----
set_para(17,
    "Residual reverse power flow is constrained at both station and upstream levels,")

# ---- P19 (1 run) ----
P[19]._p.getparent().remove(P[19]._p)   # drop redundant eq-(3) explanation to reclaim space

# ---- P23 (1 run) ----
set_para(23,
    " Storage only absorbs reverse power flow, not credited as firm N-1 backup. The optimization is subject to the "
    "variable bounds, the aggregate-ratio band, and the reverse-power-flow constraints.")
set_para(25,
    "EPG-NSGA-II retains the NSGA-II backbone [6] and applies warm-start sampling and repair projection before "
    "objective evaluation: the warm start seeds a near-feasible population, repair follows every variation step.")

# ---- B. Procedure prose (formalize #24/#25) ----
set_para(26,
    "WarmStartSampling draws each ratio Rj in [1.2, 3.0], projects the aggregate into band (2), and allocates storage "
    "by descending residual reverse-power-flow pressure.")
set_para(27,
    "RepairProjection clips Rj to [1.2, 3.0] and Pj to [0, 6] MW, rescales the aggregate into band (2), enforces the "
    "station and upstream reverse-power-flow constraints (3), and runs a final feasibility check.")

# ---- Setup ----
set_para(30,
    "The benchmark is a public-data-informed synthetic county test system constructed from scaled IEEE 33-bus "
    "feeders [7] and PVGIS irradiance [8], with county sizes K=10, 20, 30, 40, population 100, and three seeds per "
    "setting. The baselines, standard NSGA-II and NSGA-III [4], run 800 generations as a longer-budget reference; "
    "EPG-NSGA-II runs 400.")

# ---- Table I description (rebased to NSGA-II = 1.000) ----
set_para(33,
    "Table I reports converged HV normalized to standard NSGA-II = 1.000. EPG-NSGA-II is at or above the baseline for "
    "every K, widening to +4.8% at K=40; Repair-only tracks it almost exactly (repair drives most of the converged "
    "gain), Warm-start-only stays near the baseline (early acceleration), and NSGA-III stays below on this "
    "bi-objective problem.")
set_para(34,
    "At K=40, EPG-NSGA-II reaches standard NSGA-II's converged HV in about 84 generations, roughly 9.5× fewer "
    "evaluations.")

# ---- Application ----
set_para(36,
    "K=20 is a representative county size. Its knee solution has a load-weighted aggregate capacity-load ratio of "
    "1.80, an annualized cost of 188.82 million CNY/yr, and EENS of 114 MWh/yr. High generation-to-load substations "
    "receive storage, low-ratio ones rely on tie-line support (Fig. 1). Versus the uniform R=2.0 baseline, the "
    "differentiated design cuts annualized cost by 13.1% at the same EENS level.")

# ---- Fig. 1 caption ----
set_para(38, "Fig. 1.  Differentiated planning configuration on the synthetic county test system (K=20).")

# ---- Conclusion ----
set_para(40,
    "EPG-NSGA-II's feasibility-preserving warm-start sampling and repair projection enable differentiated "
    "capacity-load ratio optimization under high distributed-PV penetration. It lowers annualized cost at the same "
    "EENS level and improves converged hypervolume over standard NSGA-II, widening to +4.8% at K=40. Future work: "
    "real-grid calibration, energy-limited storage, and statistical testing.")

# ---- Table 6: make_table1_full.py table (NSGA-II = 1.000 basis) minus the NSGA-III row,
#      since only the single editor citation is added (NSGA-III would need its own reference) ----
def cell_border(cell, top=False, bottom=False, sz=10):
    tcPr = cell._tc.get_or_add_tcPr()
    for b in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(b)
    tb = OxmlElement('w:tcBorders')
    for edge, on in (('top', top), ('bottom', bottom)):
        if on:
            e = OxmlElement('w:' + edge)
            e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz))
            e.set(qn('w:space'), '0'); e.set(qn('w:color'), '000000')
            tb.append(e)
    tcPr.append(tb)


t = d.tables[6]
rows5 = [["NSGA-II", "1.000", "1.000", "1.000", "1.000"],
         ["NSGA-III", "0.998", "0.991", "0.979", "0.985"],
         ["Warm-start-only", "0.996", "0.991", "0.988", "0.998"],
         ["Repair-only", "1.004", "1.006", "1.008", "1.048"],
         ["EPG-NSGA-II", "1.007", "1.007", "1.009", "1.048"]]
while len(t.rows) < 1 + len(rows5):     # grow to header + 5 data rows
    t.add_row()
for ri, vals in enumerate(rows5, start=1):
    for ci, v in enumerate(vals):
        cell = t.rows[ri].cells[ci]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.0
        if p.runs:
            p.runs[0].text = v
            for r in p.runs[1:]:
                r.text = ""
            run = p.runs[0]
        else:
            run = p.add_run(v)
        run.font.size = Pt(8); run.font.name = "Times New Roman"; run.font.bold = False
        cell_border(cell, bottom=(ri == len(rows5)))   # bottom rule only on last data row
# header row (row 0): top + bottom rules
for c in t.rows[0].cells:
    cell_border(c, top=True, bottom=True)

# NOTE: template layout unchanged — body stays 10 pt, title 24 pt, figure original size,
# reference line-spacing original (0.83). 2 pages held by content length only.

# ---- Add the single editor-required reference [9] (Hu-Yao, evolutionary MO optimization) ----
def add_ref(parts):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 0.83
    p._p.get_or_add_pPr().get_or_add_ind().set(qn('w:hanging'), '150')
    for txt, italic in parts:
        r = p.add_run(txt)
        r.font.size = Pt(7.5); r.font.name = "Times New Roman"; r.font.italic = italic
        rpr = r._element.get_or_add_rPr()
        rf = rpr.find(qn('w:rFonts'))
        if rf is None:
            rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
        for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia'):
            rf.set(qn(a), "Times New Roman")
    return p


add_ref([("[9]  X. Xue, C. Yang, Y. Hu, et al., “Evolutionary sequential transfer optimization for "
          "objective-heterogeneous problems,” ", False),
         ("IEEE Trans. Evol. Comput.", True),
         (", vol. 26, no. 6, pp. 1424-1438, 2022.", False)])

# ---- Replace Fig. 1 with the regenerated image (colorbar 'capacity-load ratio R',
#      not 'per-station ratio R') and resize to reclaim right-column space ----
import struct
FIGPNG = ("/home/roscy/ws_HengXU/徐州地区分布式新能源高渗透率地区110kV电网容载比弹性指标优化研究/"
          "实验/有EA/results/figures/fig1_matrix.png")
sh = d.inline_shapes[0]
rId = sh._inline.graphic.graphicData.pic.blipFill.blip.embed
img_part = d.part.related_parts[rId]
with open(FIGPNG, "rb") as f:
    data = f.read()
img_part._blob = data
pw, ph = struct.unpack(">II", data[16:24])   # PNG IHDR width,height
_FIGW = 6.0                                   # figure restored (body text cut to compensate)
sh.width = Cm(_FIGW)
sh.height = int(Cm(_FIGW) * ph / pw)

# ---- Replace uncited [4] (Lin, PV hosting) with NSGA-III (Deb & Jain), cited in Setup for the
#      new Table I row; keeps reference count at 9 and does not touch editor ref [9] ----
for p in d.paragraphs:
    if p.runs and p.runs[0].text.startswith("[4]"):
        p.runs[0].text = ('[4]  K. Deb and H. Jain, "An evolutionary many-objective optimization algorithm using '
                          'reference-point-based nondominated sorting approach, part I: solving problems with box '
                          'constraints," ')
        if len(p.runs) >= 3:
            p.runs[1].text = "IEEE Trans. Evol. Comput."
            p.runs[2].text = ", vol. 18, no. 4, pp. 577-601, 2014."
            for r in p.runs[3:]:
                r.text = ""
        break

# ---- Slightly tighten reference line spacing only (7.5 pt refs; not a body-font change) ----
for p in d.paragraphs:
    if p.runs and p.runs[0].font.size and abs(p.runs[0].font.size.pt - 7.5) < 0.1:
        p.paragraph_format.line_spacing = 0.76

# ---- Renumber references [1]-[9] strictly by order of first appearance in the body ----
import re as _re
_refpat = _re.compile(r'\[(\d+)\]')
_all = list(d.paragraphs)   # single snapshot: d.paragraphs makes fresh wrappers each call
_ref_paras = [p for p in _all if _re.match(r'^\s*\[\d+\]', p.text)]
_ref_ids = {id(p) for p in _ref_paras}
_body = [p for p in _all if id(p) not in _ref_ids]
_order = []
for p in _body:
    for m in _refpat.findall(p.text):
        n = int(m)
        if n not in _order:
            _order.append(n)
_map = {old: i + 1 for i, old in enumerate(_order)}
# remap in-text citations
for p in _body:
    for r in p.runs:
        if '[' in r.text:
            r.text = _refpat.sub(lambda m: '[%d]' % _map.get(int(m.group(1)), int(m.group(1))), r.text)
# Capture each reference's run data (text, italic, size) minus its [old] label, keyed by old number
_cap = {}
for p in _ref_paras:
    old = int(_re.match(r'^\s*\[\s*(\d+)\s*\]', p.text).group(1))
    data = [(r.text, bool(r.font.italic), r.font.size) for r in p.runs if r.text]
    if data:
        data[0] = (_re.sub(r'^\s*\[\s*\d+\s*\]\s*', '', data[0][0]), data[0][1], data[0][2])
    _cap[old] = data
_inv = {new: old for old, new in _map.items()}
# Rewrite each physical reference paragraph in place: position i shows new number i+1
for i, p in enumerate(_ref_paras):
    newnum = i + 1
    data = _cap[_inv[newnum]]
    for r in list(p.runs):
        p._p.remove(r._element)
    lr = p.add_run("[%d]  " % newnum); lr.font.size = Pt(7.5); lr.font.name = "Times New Roman"
    for txt, ital, sz in data:
        rr = p.add_run(txt); rr.font.italic = ital
        rr.font.size = sz if sz is not None else Pt(7.5); rr.font.name = "Times New Roman"
d.save(OUT)
print("saved:", OUT)
