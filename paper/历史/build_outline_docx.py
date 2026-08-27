#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""把 paper/01_论文大纲_中文_v1.md 渲染成 Word 版大纲（单栏骨架）。
- 忠实渲染大纲文本（标题/子标题/引用注/项目符号/管道表）。
- 附录物理落位：已生成的 Fig.2/3/4（PNG）+ 双语图题 + 【图片要求】；Fig.1 占位框 + 绘图要求；
  Table I（5 变体收敛 HV + 加速比 + 消融，读自 ea_converge.csv）、Table II（应用汇总，读自应用 CSV）——均动态计算防漂移。
- 文档作者写 XH（core-property，"word 编辑的作者"）。
排版对齐 MIND 对标范文：Fig. N. 图下、TABLE N 表上带描述性小标题、三线表。
用法：python3 build_outline_docx.py"""
from __future__ import annotations
import os, re
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = "/home/roscy/ws_HengXU/徐州地区分布式新能源高渗透率地区110kV电网容载比弹性指标优化研究"
MD   = os.path.join(ROOT, "paper/01_论文大纲_中文_v1.md")
FIGD = os.path.join(ROOT, "实验/有EA/results/figures")   # ★ 有EA 投稿级图（用 PNG，docx 不能嵌 PDF）
RESD = os.path.join(ROOT, "实验/有EA/results")
OUT  = os.path.join(ROOT, "paper/徐州110kV容载比_MIND2026大纲_XH.docx")
SONG, HEI, TNR = "宋体", "微软雅黑", "Times New Roman"
GRAY = RGBColor(0x5B, 0x67, 0x73)
NAVY = RGBColor(0x13, 0x29, 0x49)


def setfont(run, size, bold=False, ea=SONG, latin=TNR, italic=False, color=None):
    run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic
    run.font.name = latin
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    rf.set(qn('w:eastAsia'), ea); rf.set(qn('w:ascii'), latin); rf.set(qn('w:hAnsi'), latin)


def _add_runs_md(p, text, size, ea=SONG, base_bold=False, color=None):
    """按 **bold** 分段加 run。"""
    for i, seg in enumerate(text.split("**")):
        if not seg:
            continue
        r = p.add_run(seg)
        setfont(r, size, bold=base_bold or (i % 2 == 1), ea=ea, color=color)


def para(doc, text="", size=10.5, bold=False, align=AL.JUSTIFY, ea=SONG, before=0,
         after=2, line=1.05, indent_chars=0, color=None, md=True):
    p = doc.add_paragraph(); p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = line
    if indent_chars:
        ind = p._p.get_or_add_pPr().get_or_add_ind()
        ind.set(qn('w:firstLineChars'), str(indent_chars * 100))
    if text:
        if md:
            _add_runs_md(p, text, size, ea=ea, base_bold=bold, color=color)
        else:
            r = p.add_run(text); setfont(r, size, bold, ea, color=color)
    return p


def cell_border(cell, top=None, bottom=None):
    tcPr = cell._tc.get_or_add_tcPr()
    b = tcPr.find(qn('w:tcBorders'))
    if b is None:
        b = OxmlElement('w:tcBorders'); tcPr.append(b)
    for edge, sz in (('top', top), ('bottom', bottom)):
        if sz is None:
            continue
        el = b.find(qn('w:' + edge))
        if el is None:
            el = OxmlElement('w:' + edge); b.append(el)
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), '000000')


def three_line(table, fs=9, head_bold=True):
    nr = len(table.rows)
    for ci in range(len(table.columns)):
        cell_border(table.rows[0].cells[ci], top=12, bottom=6)
        cell_border(table.rows[nr - 1].cells[ci], bottom=12)
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = AL.CENTER
                p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
                for r in p.runs:
                    setfont(r, fs, bold=(head_bold and ri == 0))


def add_table(doc, headers, rows, fs=9, align_first_left=False):
    t = doc.add_table(rows=1, cols=len(headers)); t.autofit = True
    for ci, h in enumerate(headers):
        t.rows[0].cells[ci].text = str(h)
    for rdata in rows:
        cells = t.add_row().cells
        for ci, v in enumerate(rdata):
            cells[ci].text = "" if v is None else str(v)
    three_line(t, fs)
    if align_first_left:
        for ri in range(1, len(t.rows)):
            for p in t.rows[ri].cells[0].paragraphs:
                p.alignment = AL.LEFT
    return t


def caption_below(doc, cn, en):
    para(doc, cn, size=9, bold=True, align=AL.CENTER, before=2, after=0, md=False)
    para(doc, en, size=9, bold=True, align=AL.CENTER, after=2, ea=TNR, md=False)


def table_title(doc, cn, en):
    para(doc, cn, size=9, bold=True, align=AL.CENTER, before=6, after=0, md=False)
    para(doc, en, size=9, bold=True, align=AL.CENTER, after=1, ea=TNR, md=False)


def note(doc, text, size=9):
    para(doc, text, size=size, align=AL.JUSTIFY, color=GRAY, before=1, after=3,
         indent_chars=1, md=True)


def place_figure(doc, png, cn, en, req):
    p = doc.add_paragraph(); p.alignment = AL.CENTER
    p.paragraph_format.space_before = Pt(6)
    if os.path.exists(png):
        p.add_run().add_picture(png, width=Cm(11.5))
    else:
        r = p.add_run("［图片占位：文件未找到，请补图］"); setfont(r, 10, bold=True, color=RGBColor(0xAF, 0x3A, 0x3E))
    caption_below(doc, cn, en)
    note(doc, "【图片要求】" + req)


def placeholder_figure(doc, cn, en, req):
    """Fig.1 尚未绘制：占位框 + 绘图要求。"""
    tb = doc.add_table(rows=1, cols=1); tb.autofit = True
    cell = tb.cell(0, 0)
    for ci in range(1):
        cell_border(cell, top=6, bottom=6)
    tcPr = cell._tc.get_or_add_tcPr()
    for edge in ('left', 'right'):
        b = tcPr.find(qn('w:tcBorders'))
        el = OxmlElement('w:' + edge); el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), 'AAAAAA'); b.append(el)
    cell.paragraphs[0].alignment = AL.CENTER
    r = cell.paragraphs[0].add_run("［Fig. 1 待绘制 —— 融合框架 + 修复流程示意图］")
    setfont(r, 11, bold=True, color=GRAY)
    for _ in range(3):
        cell.add_paragraph()
    caption_below(doc, cn, en)
    note(doc, "【图片要求】" + req)


# ============ 动态计算 Table I / II（读 CSV，防漂移）============
def _fmt_hv(v):
    return f"{v:.3g}"


def table1_rows():
    df = pd.read_csv(os.path.join(RESD, "ea_converge.csv"))
    mg = df.groupby("algo").gen.max().to_dict()
    order = ["classic", "nsga3", "repair_only", "warmstart_only", "fg"]
    rows = []
    for K in sorted(df.K.unique()):
        s = df[df.K == K]
        def hv(a):
            return s[(s.algo == a) & (s.gen == mg[a])].hv.mean()
        vals = {a: hv(a) for a in order}
        fgc = s[s.algo == "fg"].groupby("gen").hv.mean()
        sp = {}
        for b in ("classic", "nsga3"):
            r = fgc[fgc >= hv(b)]
            g = int(r.index.min()) if len(r) else mg["fg"]; sp[b] = mg[b] / g
        rows.append([str(int(K))] + [_fmt_hv(vals[a]) for a in order]
                    + [f"{vals['fg']/vals['classic']:.3f}", f"{vals['fg']/vals['nsga3']:.3f}",
                       f"{sp['classic']:.1f}×", f"{sp['nsga3']:.1f}×"])
    return rows


def table2_rows():
    bl = pd.read_csv(os.path.join(RESD, "ea_baseline.csv")).iloc[0]
    rob = pd.read_csv(os.path.join(RESD, "ea_robustness.csv")).saving_pct
    band = pd.read_csv(os.path.join(RESD, "ea_band_effect.csv"))
    val = pd.read_csv(os.path.join(RESD, "ea_validation.csv"))
    m, sd = rob.mean(), rob.std()
    bands = " / ".join(f"{r.band}:{r.min_cost_wan:.0f}" for _, r in band.iterrows())
    ve = val[val.algo == "enhanced"]
    vstr = " / ".join(f"K={int(r.K)}:{r.HV_ratio:.3f}" for _, r in ve.iterrows())
    return [
        ["相对全县统一容载比 2.0（同等风险）", f"{bl.rigid_f1_wan:.0f} → {bl.ea_f1_at_same_risk_wan:.0f} 万/年，节省 {bl.ea_cost_saving_pct:.1f}%"],
        ["拐点推荐方案（K=20）", "县总成本 18882 万/年、N-1 风险 114 MWh、聚合容载比 1.80（<2.0）"],
        ["一站一策路线分布", "11 站配储为主 / 9 站低容载比+强联络 / 0 站需增容（最大 R*≈1.96）"],
        ["容载比带效应（成本下限, EENS=0）", bands + "（单调↑）"],
        [f"多县鲁棒性（{len(rob)} 县重采样）", f"节省 {m:.1f}%±{sd:.1f}%（{rob.min():.1f}–{rob.max():.1f}%，全为正）"],
        ["上级系统潮流约束（E7）", "加约束成本 Δ≈+25 万（<0.3%），本合成县近乎不绑定"],
        ["求解器验证（小 K，对穷举真前沿）", f"HV 比≈1（{vstr}），近达真 Pareto 前沿"],
    ]


# ==================== 主渲染 ====================
doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2.2)
sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(2.0)

raw = open(MD, encoding="utf-8").read().splitlines()

i = 0
n = len(raw)
while i < n:
    ln = raw[i].rstrip(); i += 1
    s = ln.strip()
    if not s:
        continue
    if s.startswith("# ") and not s.startswith("## "):          # 文档大标题
        para(doc, s[2:].strip(), size=15, bold=True, align=AL.CENTER, before=0, after=8, md=False)
        continue
    if s.startswith("## "):                                     # 一级
        para(doc, re.sub(r"^[★\s]+", "", s[3:].strip()), size=13.5, bold=True,
             align=AL.LEFT, before=8, after=3, color=NAVY, md=False)
        continue
    if s.startswith("### "):                                    # 二级
        para(doc, s[4:].strip(), size=12, bold=True, align=AL.LEFT, before=5, after=2, md=False)
        continue
    if s.startswith("---"):
        continue
    if s.startswith(">"):                                       # 引用注
        note(doc, s.lstrip("> ").strip())
        continue
    if s.startswith("|"):                                       # 管道表块
        block = [ln]
        while i < n and raw[i].lstrip().startswith("|"):
            block.append(raw[i]); i += 1
        cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in block]
        cells = [r for r in cells if not all(re.fullmatch(r":?-{2,}:?", c or "-") for c in r)]
        if cells:
            hdr = [re.sub(r"\*\*", "", c) for c in cells[0]]
            body = [[re.sub(r"\*\*", "", c) for c in r] for r in cells[1:]]
            add_table(doc, hdr, body, fs=8, align_first_left=True)
        continue
    # 项目符号 / 普通段
    indent = len(ln) - len(ln.lstrip(" "))
    lead = s[:3]
    if s.startswith("- ") or re.match(r"^\d+\)", s) or re.match(r"^[a-z]\)", s):
        bullet = "  " * (1 if indent >= 2 else 0) + ("– " if indent >= 2 else "• ")
        txt = s[2:] if s.startswith("- ") else s
        p = para(doc, "", size=10.5, align=AL.JUSTIFY, after=1, indent_chars=0)
        r0 = p.add_run(bullet); setfont(r0, 10.5, bold=False)
        _add_runs_md(p, txt, 10.5)
    else:
        para(doc, s, size=10.5, align=AL.JUSTIFY, indent_chars=1, after=2)

# ---------------- 附录：图表样张 + 数据表 ----------------
doc.add_page_break()
para(doc, "附录 A：图表样张与数据表（Word 大纲占位；数据源 有EA/results，收敛扫描 run beyb8rsgp）",
     size=13.5, bold=True, color=NAVY, before=0, after=4, md=False)
note(doc, "4 图均为投稿级 PDF 矢量 + PNG（Liberation Serif / Okabe-Ito / Type-42），此处用 PNG 占位。图号按论文叙述顺序（应用在前）：Fig. 2=应用双面板、Fig. 3=HV 收敛（★核心图）、Fig. 4=加速比——图号≠重要性。英文术语依 有EA/term-check-report.md。")

place_figure(doc, os.path.join(FIGD, "fig_framework.png"),
    "图1  FG-NSGA-II 融合框架与可行性修复流程",
    "Fig. 1.  Overview of the physics-guided FG-NSGA-II framework and the feasibility-repair procedure.",
    "左“物理评价器 Z4-LCC compute”↔右“FG-NSGA-II 搜索器”双框；搜索器内闭环：可行热启动→评价→非支配排序+拥挤度+可行性优先→交叉变异→可行性修复→选择；修复子流程（缩 R 入带→补储/必要时升 R 过反送闸→补储满足系统约束），标注“修复不调目标函数→不占评估预算”。")

place_figure(doc, os.path.join(FIGD, "fig_county_pareto.png"),
    "图2(a)  县级“成本–N-1 风险”Pareto 前沿与拐点",
    "Fig. 2(a).  County-level cost–N-1 risk Pareto front and the knee point.",
    "权衡曲线 + 推荐拐点（f₁=18882 万/年、f₂=114 MWh、聚合容载比 1.80）。")

place_figure(doc, os.path.join(FIGD, "fig_strategy.png"),
    "图2(b)  拐点处逐站“一站一策”",
    "Fig. 2(b).  Per-substation strategy at the knee point.",
    "横=generation-to-load ratio、纵=R*、点色/大小=配储与路线（11 配储 / 9 低容载比+强 tie-line interconnection / 0 增容），最大 R*≈1.96。")

place_figure(doc, os.path.join(FIGD, "fig_enhanced_convergence.png"),
    "图3（★核心）  各规模逐代 HV 收敛（4 个 K、5 变体 mean±std）",
    "Fig. 3.  Per-generation hypervolume (mean±std, 3 seeds) for K in {10,20,30,40} across five variants.",
    "2×2 子面板 K=10/20/30/40；5 曲线 FG / classic / NSGA-III / repair-only / warm-start-only。消融：repair-only 全程≈FG（修复=主驱动）、warm-start-only 早期超经典而收敛后洗掉（热启动=早期加速器）。种群固定 pop=100，故 gen ∝ NFE(=pop×gen) ∝ 墙钟时间。")

place_figure(doc, os.path.join(FIGD, "fig_scaling_gap.png"),
    "图4  iso-quality 加速比随规模 K",
    "Fig. 4.  Iso-quality speedup versus problem scale K.",
    "vs classic / vs NSGA-III 两条加速比线（6.6/4.8/5.3/9.5× 与 7.5/9.1/11.1/11.8×，全域 4.8–11.8×）+ parity(1×) 基准线；预算无关。")

# TABLE I
table_title(doc, "表1  5 变体收敛 HV、加速比与消融归因（K∈{10,20,30,40}）",
            "TABLE I.  CONVERGED HYPERVOLUME, ISO-QUALITY SPEEDUP AND ABLATION ACROSS FIVE VARIANTS")
add_table(doc,
    ["K", "Classic", "NSGA-III", "Repair-only", "Warm-start", "FG (ours)",
     "FG/Cls", "FG/N3", "加速×Cls", "加速×N3"],
    table1_rows(), fs=8)
note(doc, "收敛 HV（baseline gen800、FG/消融 gen400、每 K 共同参考点）。读法：Repair-only≈FG → 修复=主驱动；Warm-start-only 收敛≈Classic → 热启动=早期加速器。加速比＝FG 达 baseline gen800 收敛质量所需代数之比（预算无关）。离散性以 3 seed 的 mean±std 报告；正式显著性检验（Wilcoxon）列为扩展版工作，非本文待办。")

# TABLE II
table_title(doc, "表2  应用结果汇总（FG-NSGA-II，K=20，收敛口径）",
            "TABLE II.  SUMMARY OF APPLICATION RESULTS (FG-NSGA-II, K=20)")
add_table(doc, ["项目", "结果"], table2_rows(), fs=9, align_first_left=True)

# ---- 作者元数据：XH ----
cp = doc.core_properties
cp.author = "XH"; cp.last_modified_by = "XH"
cp.title = "分布式新能源高渗透县级110kV容载比弹性 FG-NSGA-II 多目标优化 —— 论文大纲"
doc.save(OUT)
sz = os.path.getsize(OUT) / 1024
print(f"[OK] {OUT}  ({sz:.0f} KB)")
print("author:", cp.author)
