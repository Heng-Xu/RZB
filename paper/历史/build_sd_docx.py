#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""按《山东电力技术》论文写作模板把初稿灌成 docx。
模板规范来源：paper/山东电力技术模版/论文写作模板.doc（经提取）：
  A4；正文双栏、题名/摘要单栏；正文宋体小四(12)+数字Times New Roman；题名宋体15粗体居中；
  一级标题宋体14粗体、二级宋体12粗体；作者/单位/摘要/中图/参考文献五号(10.5)；
  图题(下)/表题(上)宋体小五(9)粗体居中、中英双语；公式右对齐编号(n)；三线表。
正文取自 paper/01_论文初稿_v1.md；图取 实验/results/figures；表取 实验/results/tables。
"""
import re, os
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from docx.enum.text import WD_TAB_ALIGNMENT
import subprocess, zipfile
from lxml import etree

ROOT = "/home/roscy/ws_HengXU/徐州地区分布式新能源高渗透率地区110kV电网容载比弹性指标优化研究"
MD   = os.path.join(ROOT, "paper/01_论文初稿_v1.md")
FIGD = os.path.join(ROOT, "实验/results/figures")
TABD = os.path.join(ROOT, "实验/results/tables")
OUT  = os.path.join(ROOT, "paper/徐州110kV容载比_山东电力技术_初稿.docx")

SONG, TNR = "宋体", "Times New Roman"

def setfont(run, size, bold=False, ea=SONG, latin=TNR, italic=False):
    run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    rf.set(qn('w:eastAsia'), ea); rf.set(qn('w:ascii'), latin); rf.set(qn('w:hAnsi'), latin)

def para(doc, text, size=12, bold=False, align=AL.JUSTIFY, ea=SONG, latin=TNR,
         indent_chars=0, before=0, after=0, line=None, italic=False):
    p = doc.add_paragraph(); p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    if line: pf.line_spacing = line
    if indent_chars:
        ind = p._p.get_or_add_pPr().get_or_add_ind()
        ind.set(qn('w:firstLineChars'), str(indent_chars*100))
    if text:
        r = p.add_run(text); setfont(r, size, bold, ea, latin, italic)
    return p

def labeled(doc, label, body, size=10.5, align=AL.JUSTIFY):
    p = doc.add_paragraph(); p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(label); setfont(r1, size, True)
    r2 = p.add_run(body); setfont(r2, size, False)
    return p

def set_cols(section, num, space=425):
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols'); sectPr.append(cols)
    cols.set(qn('w:num'), str(num))
    if num > 1: cols.set(qn('w:space'), str(space))
    else:
        if cols.get(qn('w:space')): del cols.attrib[qn('w:space')]

def cell_border(cell, top=None, bottom=None):
    tcPr = cell._tc.get_or_add_tcPr()
    b = tcPr.find(qn('w:tcBorders'))
    if b is None:
        b = OxmlElement('w:tcBorders'); tcPr.append(b)
    for edge, sz in (('top', top), ('bottom', bottom)):
        if sz is None: continue
        el = b.find(qn('w:'+edge))
        if el is None:
            el = OxmlElement('w:'+edge); b.append(el)
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), '000000')

def three_line(table, fs=9):
    """三线表：表首顶线、表头下线、表底线；单元格宋体小五。
    长表跨页时表头行自动重复（w:tblHeader），且各行不被分页截断（w:cantSplit）。"""
    nr = len(table.rows)
    for ci in range(len(table.columns)):
        cell_border(table.rows[0].cells[ci], top=12, bottom=6)
        cell_border(table.rows[nr-1].cells[ci], bottom=12)
    # 表头行跨页重复
    trPr0 = table.rows[0]._tr.get_or_add_trPr()
    th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), 'true'); trPr0.append(th)
    for ri, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        cs = OxmlElement('w:cantSplit'); cs.set(qn('w:val'), 'true'); trPr.append(cs)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = AL.CENTER
                for r in p.runs:
                    setfont(r, fs, bold=(ri == 0))
                if not p.runs:
                    pass

def add_table(doc, headers, rows, fs=9, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.autofit = True
    for ci, h in enumerate(headers):
        t.rows[0].cells[ci].text = str(h)
    for rdata in rows:
        cells = t.add_row().cells
        for ci, v in enumerate(rdata):
            cells[ci].text = "" if v is None else str(v)
    three_line(t, fs)
    return t

def caption(doc, cn, en):
    para(doc, cn, size=9, bold=True, align=AL.CENTER, after=0, before=3)
    para(doc, en, size=9, bold=True, align=AL.CENTER, after=4, latin=TNR)

def add_figure(doc, num):
    fmap = {1:"fig1_dilemma.png",2:"fig2_framework.png",3:"fig5_lcc_comparison.png",
            4:"fig10_reverse_capacity.png",5:"fig3_decision_matrix.png",6:"fig9_growth_sensitivity.png",
            7:"fig8_regional.png"}
    cap = {1:("图1  刚性容载比的双向压力两难","Fig.1  The bidirectional-pressure dilemma of rigid capacity-load ratio"),
           2:("图2  多维分类决策框架（含反向承载力校核准入）","Fig.2  Multi-dimensional classification framework (with reverse hosting-capacity gate)"),
           3:("图3  决策方法与刚性2.0年化成本对比","Fig.3  Annualized cost: decision method vs. rigid 2.0"),
           4:("图4  反向承载力可行域","Fig.4  Reverse hosting-capacity feasible region"),
           5:("图5  容载比决策矩阵","Fig.5  Capacity-load ratio decision matrix"),
           6:("图6  负荷增长与N-1过载敏感性","Fig.6  Sensitivity to load growth and N-1 overload"),
           7:("图7  反送时长决定治理路线（区域算例）","Fig.7  Reverse-flow duration determines the treatment route")}
    p = doc.add_paragraph(); p.alignment = AL.CENTER; p.paragraph_format.space_before = Pt(4)
    p.add_run().add_picture(os.path.join(FIGD, fmap[num]), width=Cm(7.6))
    caption(doc, *cap[num])

# ---------- 公式（占位可读文本；投稿用公式编辑器最终排版） ----------
FORMULA = {
 1:"R_s = (Σ S_T,i) / P_max",
 2:"λ = P_PV / P_max ,   η = E_PV / E_load",
 3:"C_ann = (C_cap − P_sv)·CRF(r,y) + C_OM + C_loss + C_cur + C_rel",
 4:"CRF(r,y) = r(1+r)^y / [(1+r)^y − 1]",
 5:"C_loss^fd = c_e · Σ_j (P_j² R_j / U²)·(α_f t_f + α_r t_r)",
 6:"C_loss^tr = c_e·[ P_0·8760 + P_k β_f² t_f + α_r P_k β_r² t_r ]",
 7:"β = S_flow / S_T",
 8:"P_ov = max(0, P_PV·k_pv − P_min − P_st − γ S_T)",
 9:"ΔS = max(0, S_pk(1+g)^n − [ k_ol·(S_T/m)(m−1) + cd·S_pk ])",
 10:"C_rel = ΔS·cosφ·H_eq·V_oll",
 11:"R_s < 2[(1+g)^n − cd] / (k_ol·cosφ)",
 12:"0.9 P_PV − 0.3 P_max ≤ γ R_s P_max,  γ=0.85",
 13:"R_s* = argmin (s∈R̃)  C_ann(s | λ, η, cd, t_r)",
}
# LaTeX 源（用于经 pandoc 生成 Word 原生公式对象 OMML；\tag 由编号单独处理）
FORMULA_LATEX = {
 1: r"R_s = \dfrac{\sum_i S_{T,i}}{P_{\max}}",
 2: r"\lambda=\dfrac{P_{PV}}{P_{\max}},\quad \eta=\dfrac{E_{PV}}{E_{load}}",
 3: r"C_{ann}=(C_{cap}-P_{sv})\,CRF(r,y)+C_{OM}+C_{loss}+C_{cur}+C_{rel}",
 4: r"CRF(r,y)=\dfrac{r(1+r)^{y}}{(1+r)^{y}-1}",
 5: r"C_{loss}^{fd}=c_e\sum_{j}\dfrac{P_{j}^{2}R_{j}}{U^{2}}(\alpha_{f}t_{f}+\alpha_{r}t_{r})",
 6: r"C_{loss}^{tr}=c_e\left[P_0\cdot 8760+P_k\beta_f^{2}t_f+\alpha_r P_k\beta_r^{2}t_r\right]",
 7: r"\beta=\dfrac{S_{flow}}{S_{T}}",
 8: r"P_{ov}=\max(0,\;P_{PV}k_{pv}-P_{\min}-P_{st}-\gamma S_T)",
 9: r"\Delta S=\max\!\left(0,\;S_{pk}(1+g)^{n}-\left[k_{ol}\tfrac{S_T}{m}(m-1)+cd\,S_{pk}\right]\right)",
 10: r"C_{rel}=\Delta S\cdot\cos\varphi\cdot H_{eq}\cdot V_{oll}",
 11: r"R_s<\dfrac{2\left[(1+g)^{n}-cd\right]}{k_{ol}\cos\varphi}",
 12: r"0.9\,P_{PV}-0.3\,P_{\max}\le \gamma R_s P_{\max},\quad \gamma=0.85",
 13: r"R_s^{*}=\arg\min_{s\in\tilde{\mathcal{R}}}\; C_{ann}(s\,|\,\lambda,\eta,cd,t_r)",
}
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
def build_omaths():
    """经 pandoc 把 12 个 LaTeX 公式批量转为 OMML（Word 原生公式对象）。"""
    try:
        md = "\n\n".join("$$%s$$" % FORMULA_LATEX[i] for i in range(1, 14))
        open("/tmp/_formulas.md", "w", encoding="utf-8").write(md)
        subprocess.run(["pandoc", "/tmp/_formulas.md", "-o", "/tmp/_formulas.docx"],
                       check=True, capture_output=True)
        root = etree.fromstring(zipfile.ZipFile("/tmp/_formulas.docx").read("word/document.xml"))
        oms = root.findall(".//{%s}oMath" % M_NS)
        return {i + 1: oms[i] for i in range(len(oms))}
    except Exception as e:
        print("[warn] OMML 生成失败，回退为可读文本：", e); return {}
OMATHS = build_omaths()

def add_formula(doc, num):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Cm(3.9), WD_TAB_ALIGNMENT.CENTER)
    tabs.add_tab_stop(Cm(8.0), WD_TAB_ALIGNMENT.RIGHT)
    p.add_run("\t")
    om = OMATHS.get(num)
    if om is not None:
        p._p.append(parse_xml(etree.tostring(om).decode()))   # 原生 Word 公式对象
    else:
        r = p.add_run(FORMULA[num]); setfont(r, 12, italic=True)
    rn = p.add_run("\t（%d）" % num); setfont(rn, 12)

# ---------- 行内 LaTeX/markdown 轻清洗 ----------
def clean_inline(s):
    s = s.replace("**", "")
    # 注意：长键须在其前缀键之前，避免 \subseteq 被 \subset 截断等
    reps = {r"\subseteq":"⊆", r"\subset":"⊂", r"\notin":"∉", r"\in":"∈",
            r"\mathcal":"", r"\mathbb":"", r"\mathrm":"",
            r"\lambda":"λ", r"\eta":"η", r"\gamma":"γ", r"\beta":"β", r"\alpha":"α",
            r"\varphi":"φ", r"\Delta":"Δ", r"\sum":"Σ", r"\times":"×", r"\cdot":"·",
            r"\cup":"∪", r"\cap":"∩", r"\to":"→", r"\neq":"≠", r"\pm":"±",
            r"\le":"≤", r"\ge":"≥", r"\approx":"≈", r"\,":"", r"\;":" ",
            r"\max":"max", r"\min":"min", r"\arg":"arg", r"\cos":"cos", r"\sin":"sin",
            r"\%":"%", r"\\":""}
    # 去 $...$ 包裹（保护 \{ \} 不被花括号清理误删）
    def fix(m):
        x = m.group(1)
        x = x.replace(r"\{", "").replace(r"\}", "")
        for k, v in reps.items(): x = x.replace(k, v)
        x = re.sub(r"\^\{([^}]*)\}", r"^\1", x); x = re.sub(r"_\{([^}]*)\}", r"_\1", x)
        x = x.replace("{","").replace("}","")
        x = x.replace("", "{").replace("", "}")
        return x
    s = re.sub(r"\$([^$]*)\$", fix, s)
    # 兜底：清理 $ 外残留的转义集合括号与常见命令
    s = s.replace(r"\{", "{").replace(r"\}", "}")
    for k, v in reps.items(): s = s.replace(k, v)
    return s.strip()

# ================= 构建文档 =================
doc = Document()
# 页面：A4 + 模板页边距（首节）
s0 = doc.sections[0]
s0.page_width, s0.page_height = Cm(21.0), Cm(29.7)
for s_set in [s0]:
    s_set.left_margin = Emu(720090); s_set.right_margin = Emu(720090)
    s_set.top_margin = Emu(895350); s_set.bottom_margin = Emu(539750)

# ---- 题名块（单栏） ----
para(doc, "高分布式新能源渗透率地区110 kV电网容载比“一片一策”弹性决策方法",
     size=15, bold=True, align=AL.CENTER, after=6)
para(doc, "许彦*，袁瑞，童伟，等", size=10.5, align=AL.CENTER, after=0)
para(doc, "（国网江苏电力设计咨询有限公司徐州勘测设计分公司，江苏 徐州　221005）",
     size=10.5, align=AL.CENTER, after=4)
labeled(doc, "摘要：",
   "随着分布式新能源高比例接入，传统以刚性单值容载比为核心的110 kV电网规划在投资效率与反送消纳、供电安全之间日益失衡。针对反送电能损耗未进入全寿命周期成本（LCC）结构、缺乏面向高压配电层的多维分类决策等问题，建立计及双向潮流电能损耗的年费用法LCC模型，将主变空载/负载损耗、反送越限弃光惩罚与N-1缺供可靠性成本统一纳入成本结构（容载比按现行导则DL/T 5729—2023净负荷口径定义）；以候选容载比逐档取最小年化成本为准则，构建（源荷比×电量渗透率×联络度×反送时长）分类决策方法，形成“一片一策”弹性容载比建议。基于IEEE 33节点系统与PVGIS实测气象的方法演示表明：刚性单值容载比2.0（处于现行导则1.5~2.0区间上限）系统性次优，最优容载比是片区特征的函数，决策方法相对刚性2.0节省年化成本约1.7%~4.7%；容载比弹性呈多维条件性，其中“刚性单值次优、需差异化整定”为稳健结论，方向与幅度随负荷增长、N-1短时过载能力、反送时长等假设变化。")
labeled(doc, "关键词：", "分布式新能源；容载比；弹性指标；全寿命周期成本；反送电；一片一策")
labeled(doc, "中图分类号：", "TM715　　文献标志码：A")
# 英文块
para(doc, "A Case-specific Elastic Decision Method for Capacity-load Ratio of 110 kV Grids with High-penetration Distributed Renewables",
     size=11, bold=True, align=AL.CENTER, before=6, after=2, ea=SONG, latin=TNR)
para(doc, "XU Yan*, YUAN Rui, TONG Wei, et al.", size=10.5, align=AL.CENTER, latin=TNR, after=0)
para(doc, "(Xuzhou Surveying and Design Branch, State Grid Jiangsu Electric Power Design and Consulting Co., Ltd., Xuzhou 221005, China)",
     size=10.5, align=AL.CENTER, latin=TNR, italic=True, after=4)
labeled(doc, "Abstract: ",
   "With the high-penetration integration of distributed renewables, conventional 110 kV grid planning centered on a rigid single-valued capacity-load ratio (CLR) is increasingly imbalanced between investment efficiency and reverse-power accommodation and supply security. To address that reverse-power energy losses are absent from the life-cycle cost (LCC) structure and that a multi-dimensional classification decision method for the high-voltage distribution layer is lacking, an annual-cost LCC model accounting for bidirectional power-flow losses is established, which unifies no-load/load losses of main transformers, reverse-overlimit PV-curtailment penalty and N-1 energy-not-supplied reliability cost. Using the minimum annualized cost over candidate CLRs as the criterion, a classification decision method along (source-load ratio × energy penetration × interconnection degree × reverse-flow duration) is built to yield case-specific elastic CLR recommendations. A method demonstration on the IEEE 33-bus system with PVGIS measured meteorology shows that the rigid single value 2.0 is systematically suboptimal and the optimal CLR is a function of zonal features, with the decision method saving about 1.7%-4.7% of annualized cost relative to rigid 2.0. The CLR elasticity is multi-dimensionally conditional: the robust conclusion is that a single rigid value is suboptimal and differentiated tuning is required, whereas the direction and magnitude depend on assumptions of load growth, N-1 short-time overload capability and reverse-flow duration.",
   align=AL.JUSTIFY)
labeled(doc, "Key words: ",
   "distributed renewables; capacity-load ratio; elasticity index; life-cycle cost; reverse power flow; case-specific strategy")

# ---- 切换为双栏正文 ----
body = doc.add_section(WD_SECTION.CONTINUOUS)
body.page_width, body.page_height = Cm(21.0), Cm(29.7)
body.left_margin = Emu(720090); body.right_margin = Emu(720090)
body.top_margin = Emu(895350); body.bottom_margin = Emu(539750)
set_cols(doc.sections[0], 1)     # 题名块单栏
set_cols(doc.sections[1], 2)     # 正文双栏

# ---- 解析 md 正文 ----
raw = open(MD, encoding="utf-8").read().splitlines()
# 取 "## 0" 到 "## 参考文献" 之间
start = next(i for i, l in enumerate(raw) if l.startswith("## 0"))
end = next(i for i, l in enumerate(raw) if l.startswith("## 参考文献"))
lines = raw[start:end]

inserted_fig, inserted_tab = set(), set()
def maybe_insert(text):
    for n in range(1, 8):
        if ("图%d" % n) in text and n not in inserted_fig:
            inserted_fig.add(n); add_figure(doc, n)
    for n in (1, 2, 3):
        if ("表%d" % n) in text and n not in inserted_tab:
            inserted_tab.add(n); add_data_table(doc, n)

def add_data_table(doc, n):
    """宽表全栏：切单栏→插表→切回双栏。"""
    import pandas as pd
    doc.add_section(WD_SECTION.CONTINUOUS)
    set_cols(doc.sections[-2], 2)
    set_cols(doc.sections[-1], 1)
    sN = doc.sections[-1]
    sN.page_width, sN.page_height = Cm(21.0), Cm(29.7)
    sN.left_margin = Emu(720090); sN.right_margin = Emu(720090)
    sN.top_margin = Emu(895350); sN.bottom_margin = Emu(539750)
    if n == 1:
        caption(doc, "表1  算例关键参数", "Table 1  Key parameters of the case study")
        add_table(doc, ["参数", "取值"], [
            ["折现率 / 计算年限", "6% / 25 a"],
            ["主变空载损耗 P0 / 负载损耗 Pk", "0.0007 / 0.0045 kW·kVA⁻¹（SZ11）"],
            ["反向限额 γ / 反向损耗权重 αr", "0.85 / 1.2"],
            ["负荷增长 g / 校核年限 n", "5% / 5 a"],
            ["N-1过载 k_ol / 失负荷价值 VOLL", "1.3 / 15 元·kWh⁻¹"],
            ["N-1等效缺供 H_eq", "90 h"],
            ["PVGIS徐州年发电量 / 容量因子", "1396 kWh·kWp⁻¹ / 15.9%"],
        ])
    elif n == 2:
        caption(doc, "表2  各类片区容载比推荐结果（IEEE33参数扫描，方法演示）", "Table 2  Recommended capacity-load ratio for each zone type (IEEE 33-bus parameter sweep, method demonstration)")
        df = pd.read_csv(os.path.join(TABD, "tab3_recommendation.csv"))
        add_table(doc, list(df.columns), df.values.tolist(), fs=8)
    elif n == 3:
        caption(doc, "表3  评价指标AHP判断矩阵（CR=0.012）", "Table 3  AHP judgment matrix of evaluation indices")
        df = pd.read_csv(os.path.join(TABD, "tab1_ahp_matrix.csv"))
        df = df.rename(columns={df.columns[0]: "指标"})
        rows = [[c if not isinstance(c, float) else round(c, 3) for c in r] for r in df.values.tolist()]
        add_table(doc, list(df.columns), rows, fs=8)
    doc.add_section(WD_SECTION.CONTINUOUS)
    set_cols(doc.sections[-2], 1)
    set_cols(doc.sections[-1], 2)
    sR = doc.sections[-1]
    sR.page_width, sR.page_height = Cm(21.0), Cm(29.7)
    sR.left_margin = Emu(720090); sR.right_margin = Emu(720090)
    sR.top_margin = Emu(895350); sR.bottom_margin = Emu(539750)

i = 0
while i < len(lines):
    ln = lines[i].rstrip(); i += 1
    if not ln.strip(): continue
    if ln.startswith("> ") or ln.startswith("---"): continue
    if ln.lstrip().startswith("|"):   # markdown 表格行（正文无，跳过）
        continue
    if ln.startswith("## "):          # 一级标题
        t = re.sub(r"（第\d+章）", "", ln[3:]).strip()
        para(doc, t, size=14, bold=True, align=AL.LEFT, before=6, after=3)
        continue
    if ln.startswith("### "):         # 二级标题
        t = ln[4:].strip()
        para(doc, t, size=12, bold=True, align=AL.LEFT, before=4, after=2)
        continue
    if ln.strip().startswith("$$"):   # 公式块（单行 $$...\tag{n}$$）
        m = re.search(r"\\tag\{(\d+)\}", ln)
        if m: add_formula(doc, int(m.group(1)))
        continue
    # 普通正文段
    txt = clean_inline(ln)
    if not txt: continue
    para(doc, txt, size=12, indent_chars=2, line=1.0, after=0)
    maybe_insert(txt)

# ---- 参考文献（双栏，五号） ----
para(doc, "参考文献", size=14, bold=True, align=AL.LEFT, before=8, after=3)
REFS = [
 "[1] DL/T 5729—2016. 配电网规划设计技术导则[S]. 北京: 中国电力出版社, 2016.",
 "[2] 国家能源局. 新型电力系统发展蓝皮书[R]. 北京: 国家能源局, 2023.",
 "[3] 宋钰, 郭玥, 王旭阳, 等. 分布式光伏接入配电网安全经济承载力评估[J]. 浙江电力, 2025, 44(11): 93-102.",
 "    SONG Yu, GUO Yue, WANG Xuyang, et al. Assessment of hosting capacity for DPV integration in distribution networks considering security and economy[J]. Zhejiang Electric Power, 2025, 44(11): 93-102.",
 "[4] 徐非非, 冯华, 覃洪培, 等. 计及不确定性的配电网分布式光伏承载能力区间分析方法[J]. 浙江电力, 2023, 42(11): 86-95.",
 "    XU Feifei, FENG Hua, QIN Hongpei, et al. Interval analysis method for DPV hosting capacity of distribution networks considering uncertainty[J]. Zhejiang Electric Power, 2023, 42(11): 86-95.",
 "[5] 王成山, 罗凤章, 肖峻, 等. 基于主变互联关系的配电系统供电能力计算方法[J]. 中国电机工程学报, 2009, 29(13): 86-91.",
 "    WANG Chengshan, LUO Fengzhang, XIAO Jun, et al. Supply capability calculation method of distribution systems based on interconnection of main transformers[J]. Proceedings of the CSEE, 2009, 29(13): 86-91.",
 "[6] 肖峻, 张婷, 张跃, 等. 基于最大供电能力的配电网规划理念与方法[J]. 中国电机工程学报, 2013, 33(10): 106-113.",
 "    XIAO Jun, ZHANG Ting, ZHANG Yue, et al. TSC-based planning idea and method for distribution networks[J]. Proceedings of the CSEE, 2013, 33(10): 106-113.",
 "[7] 肖峻, 贺国伟. 考虑双向能量交换型负荷的配电网最大供电能力模型[J]. 电力系统自动化, 2022, 46(5): 11-24.",
 "    XIAO Jun, HE Guowei. Total supply capability model of distribution networks considering bidirectional energy-exchange loads[J]. Automation of Electric Power Systems, 2022, 46(5): 11-24.",
 "[8] 荆朝霞, 江昌旭, 王宏益. 考虑N-1安全约束的220 kV片区电网最大供电能力计算[J]. 电力系统自动化, 2016, 40(19): 145-152.",
 "    JING Zhaoxia, JIANG Changxu, WANG Hongyi. Total supply capability calculation of 220 kV area power grid considering N-1 security constraint[J]. Automation of Electric Power Systems, 2016, 40(19): 145-152.",
 "[9] 刘炜彬, 梁咏秋, 李锡刚, 等. 面向精细化规划的容载比参数优化方法研究[J]. 电气工程学报, 2025, 20(1): 167-175. DOI: 10.11985/2025.01.020.",
 "    LIU Weibin, LIANG Yongqiu, LI Xigang, et al. Research on capacity-load ratio parameter optimization for fine planning[J]. Journal of Electrical Engineering, 2025, 20(1): 167-175.",
 "[10] 徐峰亮, 王克谦, 王文豪, 等. 计及运行灵活性的中压配电系统源-网-储协同扩展规划[J]. 中国电力, 2024, 57(7): 98-108.",
 "    XU Fengliang, WANG Keqian, WANG Wenhao, et al. Source-network-storage coordinated expansion planning of MV distribution systems considering operational flexibility[J]. Electric Power, 2024, 57(7): 98-108.",
 "[11] 李敬如, 白宇, 王旭阳, 等. 考虑主配协同的分布式光伏接入电网安全经济承载力计算方法[J]. 电力建设, 2025, 46(10): 113-121.",
 "    LI Jingru, BAI Yu, WANG Xuyang, et al. A calculation method for DPV safe and economic carrying capacity considering transmission-distribution coordination[J]. Electric Power Construction, 2025, 46(10): 113-121.",
 "[12] 梁志峰, 夏俊荣, 孙檬檬, 等. 数据驱动的配电网分布式光伏承载力评估技术研究[J]. 电网技术, 2020, 44(7): 2430-2439.",
 "    LIANG Zhifeng, XIA Junrong, SUN Mengmeng, et al. Data-driven assessment of distributed photovoltaic hosting capacity in distribution network[J]. Power System Technology, 2020, 44(7): 2430-2439.",
 "[13] SAATY T L. The analytic hierarchy process[M]. New York: McGraw-Hill, 1980.",
 "[14] 肖峻, 刘柔嘉, 龙梦皓. 变电容载比与下一级中压配电网络关系的量化分析[J]. 电力建设, 2015, 36(11): 45-50.",
 "    XIAO Jun, LIU Roujia, LONG Menghao. Quantitative analysis of relationship between substation capacity-to-load ratio and subordinate medium voltage distribution network[J]. Electric Power Construction, 2015, 36(11): 45-50.",
 "[15] 刘文霞, 刘其达, 邓诗语, 等. 基于容载比色带的区域配电网充裕性分析[J]. 华北电力大学学报(自然科学版), 2022, 49(2): 1-12.",
 "    LIU Wenxia, LIU Qida, DENG Shiyu, et al. Sufficiency analysis of regional distribution network based on the color band of capacity-load ratio[J]. Journal of North China Electric Power University, 2022, 49(2): 1-12.",
 "[16] 陈浩, 张焰, 俞国勤, 等. 多电压等级配电网最优容载比的计算方法[J]. 刊名待核, 2009/2010, 卷(期): 页码待核.（PDF已下载, 收稿日期2009-06-15; 刊名/卷期/页码请据原刊首页核定）",
 "    CHEN Hao, ZHANG Yan, YU Guoqin, et al. Calculation method of optimal capacity-load ratio for multi-voltage-level distribution networks[J]. (journal/vol./pp. to be verified, received 2009-06-15).",
 "[17] 李胜洪, 王家斌, 张巧霞. 湖北电网容载比问题的初步探讨[J]. 湖北电力, 2000, 24(3): 待核.（DOI:10.19308/j.hep.2000.03.016, 页码以原刊核定）",
 "    LI Shenghong, WANG Jiabin, ZHANG Qiaoxia. Preliminary discussion on the capacity-load ratio of Hubei power grid[J]. Hubei Electric Power, 2000, 24(3).",
 "[18] 栗峰, 丁杰, 周才期, 等. 新型电力系统下分布式光伏规模化并网运行关键技术探讨[J]. 电网技术, 2024, 48(1): 184-196.",
 "    LI Feng, DING Jie, ZHOU Caiqi, et al. Key technologies of large-scale grid-connected operation of distributed photovoltaic under new-type power system[J]. Power System Technology, 2024, 48(1): 184-196.",
 "[19] 黄海泉, 黄晓巍, 姜望, 等. 新型配电网分布式储能系统方案及配置研究综述[J]. 南方能源建设, 2024, 11(4): 42-53. DOI: 10.16516/j.ceec.2024.4.05.",
 "    HUANG Haiquan, HUANG Xiaowei, JIANG Wang, et al. A review of distributed energy storage system solutions and configurations for new distribution grids[J]. Southern Energy Construction, 2024, 11(4): 42-53.",
 "[20] 刘洪波, 刘珅诚, 盖雪扬, 等. 高比例新能源接入的主动配电网规划综述[J]. 发电技术, 2024, 45(1): 151-161. DOI: 10.12096/j.2096-4528.pgt.22106.",
 "    LIU Hongbo, LIU Shencheng, GAI Xueyang, et al. Overview of active distribution network planning with high proportion of new energy access[J]. Power Generation Technology, 2024, 45(1): 151-161.",
 "[21] 曹炜, 董浩洋, 李芸, 等. 分布式光伏高比例接入的国外经验及实践启示[J]. 电气传动, 2022, 52(4): 待核. DOI: 10.19457/j.1001-2095.dqcd22454.",
 "    CAO Wei, DONG Haoyang, LI Yun, et al. Foreign experience and practical enlightenment of high-proportion access of distributed photovoltaic[J]. Electric Drive, 2022, 52(4).",
 "[22] 黄凯东. 提升配电网承载水平, 探究分布式光伏消纳难题[J]. 工程与技术创新, 2025(6): 待核. DOI: 10.63887/jeti.2025.1.6.16.",
 "    HUANG Kaidong. Improving distribution network hosting capacity and exploring distributed photovoltaic accommodation[J]. Journal of Engineering and Technological Innovation, 2025(6).",
 "[23] 刘佳男, 李业行, 张董, 等. 基于风光储一体化的新型农村配电网研究[J]. 山东电力技术, 2021, 48(11): 42-49.",
 "    LIU Jianan, LI Yehang, ZHANG Dong, et al. Study on new rural distribution network based on the integration of wind, solar energy and storage[J]. Shandong Electric Power, 2021, 48(11): 42-49.",
 "[24] 龙宇, 刘晓峰, 刘怀, 等. 高比例分布式光伏接入下配电网电压有功-无功鲁棒控制[J]. 综合智慧能源, 2026, 48(1): 67-77. DOI: 10.3969/j.issn.2097-0706.2026.01.007.",
 "    LONG Yu, LIU Xiaofeng, LIU Huai, et al. Active-reactive power robust control for distribution network voltage under high-proportion distributed photovoltaic integration[J]. Integrated Intelligent Energy, 2026, 48(1): 67-77.",
 "[25] 杜剑, 欧阳俊, 李国柱, 等. 新形势下城市输电网规划方法研究[J]. 湖北电力, 2017, 41(3): 1-5. DOI: 10.19308/j.hep.2017.03.001.",
 "    DU Jian, OUYANG Jun, LI Guozhu, et al. Research on the network planning of city transmission under new situation[J]. Hubei Electric Power, 2017, 41(3): 1-5.",
 "[26] 刘克权, 汤文, 陈钊, 等. 新能源高占比下主配协同调度的挑战与提升策略[J]. 智能电网, 2026, 16(1): 7-15. DOI: 10.12677/sg.2026.161002.",
 "    LIU Kequan, TANG Wen, CHEN Zhao, et al. Challenges and improvement strategies of transmission-distribution coordinated dispatch under high proportion of new energy[J]. Smart Grid, 2026, 16(1): 7-15.",
 "[27] WANG R, JI H, LI P, et al. Multi-resource dynamic coordinated planning of flexible distribution network[J]. Nature Communications, 2024, 15: 4576.",
 "[28] 吴迪, 俞露稼, 汪若宇, 等. 分布式光伏高质量发展（山东省新型电力系统系列研究）[R]. 北京: 北京大学能源研究院, 自然资源保护协会(NRDC), 2025.",
 "[29] HAIDER R, FERRO G, ROBBA M, et al. Flattening the duck curve: a case for distributed decision making[R/OL]. arXiv:2111.06361, 2022.",
 "[30] MATTHISS B, MOMENIFARAHANI A, BINDER J. Storage placement and sizing in a distribution grid with high PV-generation[C]. 会议来源待核（PDF已下载, ZSW Baden-Württemberg）.",
]
for r in REFS:
    p = para(doc, r, size=10.5, align=AL.JUSTIFY, after=0, line=1.0, latin=TNR)

# 文档核心属性：作者写 XH（不留 pandoc/python-docx 默认值）
cp = doc.core_properties
cp.author = "XH"
cp.last_modified_by = "XH"
doc.save(OUT)
print("[OK] saved:", OUT)
print("sections:", len(doc.sections), "| figs:", sorted(inserted_fig), "| tabs:", sorted(inserted_tab))
