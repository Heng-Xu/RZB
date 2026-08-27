#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Second-round high-precision revision per 论文第二轮高精度修正说明书_供Agent执行.md.
Input = round-1 output (spec-stated input). Output = final docx.
Stage 1 here: text replacements (final [1]-[8] citation numbers), ref delete+renumber,
non-dominated scoping. Stages 2-3 (equation OMML rebuild, Algorithm 1) appended later."""
import shutil, re, subprocess, zipfile, copy
from docx import Document
from docx.shared import Pt
from lxml import etree
M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
TMP_EQ = "/tmp/claude-1000/-home-roscy-ws-HengXU-----------------110kV-------------/b58606b0-0593-4e06-b592-ccea2e68fc85/scratchpad/_eq2.docx"


def _scale_math(node, half_pts=20):
    for mr in node.iter("{%s}r" % M):
        rpr = etree.SubElement(mr, "{%s}rPr" % W)
        for t in ("sz", "szCs"):
            e = etree.SubElement(rpr, "{%s}%s" % (W, t)); e.set("{%s}val" % W, str(half_pts))
        mpr = mr.find("{%s}rPr" % M); mr.remove(rpr); mr.insert(1 if mpr is not None else 0, rpr)


def make_omathpara(latex):
    subprocess.run(["pandoc", "-f", "markdown", "-t", "docx", "-o", TMP_EQ],
                   input="$$%s$$" % latex, text=True, check=True)
    tree = etree.fromstring(zipfile.ZipFile(TMP_EQ).read("word/document.xml"))
    node = copy.deepcopy(tree.find(".//{%s}oMathPara" % M))
    _scale_math(node)
    return node


def set_eq_para(para, latex):
    for op in para._p.findall("{%s}oMathPara" % M):
        op.getparent().remove(op)
    para._p.append(make_omathpara(latex))

ROOT = "/home/roscy/ws_HengXU/徐州地区分布式新能源高渗透率地区110kV电网容载比弹性指标优化研究"
SRC = ROOT + "/paper/废/epg-nsga-ii-paper_final_compressed_revised.docx"
OUT = ROOT + "/paper/许恒-EPG-NSGA-II-终稿.docx"
shutil.copy(SRC, OUT)
d = Document(OUT)
P = d.paragraphs


def setrun0(p, text):
    """Set whole paragraph text into run0, clear the rest (pure-prose paragraphs)."""
    assert p.runs, "no runs"
    p.runs[0].text = text
    for r in p.runs[1:]:
        r.text = ""


def repl(p, old, new):
    hit = False
    for r in p.runs:
        if old in r.text:
            r.text = r.text.replace(old, new); hit = True
    return hit


# ============ Title (§4.1) ============
P[0].runs[0].text = ("Engineering-Prior-Guided NSGA-II for Differentiated Capacity-Load Ratio "
                     "Optimization in a County-Level 110-kV Grid under High Distributed PV ")
P[0].runs[1].text = "Penetration"

# ============ Authors: replaced by an IEEE 3-per-row grid (6 placeholder authors) built as STAGE 5 (end),
# so removing the old single-affiliation paragraphs P[1..5] does not shift the index-based edits below. ============

# ============ Abstract (§5) — 'final HV' per §11.2 (values verified as final-gen, mean of 3 runs) ============
P[7].runs[0].text = (
    "Abstract—High penetration of distributed photovoltaic (PV) generation can cause reverse power flow in "
    "county-level 110-kV grids and alter supply adequacy under N-1 contingencies. This study jointly optimizes "
    "differentiated capacity-load ratios (CLRs) and storage ratings by minimizing annualized cost and expected "
    "energy not supplied (EENS). Engineering-Prior-Guided NSGA-II (EPG-NSGA-II) introduces engineering-prior "
    "population seeding and feasibility repair to improve constrained search. On a public-data-informed "
    "county-level test system, a selected Pareto solution reduces annualized cost by 13.1% at the same EENS. For "
    "K=40, EPG-NSGA-II achieves a 4.8% higher mean final hypervolume than standard NSGA-II and reaches the "
    "latter's 800-generation final level in about 84 generations.")
for r in P[7].runs[1:]:
    r.text = ""

# ============ Index Terms (§6) ============
P[8].runs[1].text = ("—Capacity-load ratio, distributed photovoltaic generation, multiobjective optimization, "
                     "reverse power flow, 110-kV grid planning.")

# ============ Introduction para 1 (§7.1), final numbering ============
setrun0(P[10],
    "High penetration of distributed PV generation can produce midday net power export at 110-kV substations, "
    "causing reverse power flow and altering supply adequacy under N-1 contingencies [1]. A uniform capacity-load "
    "ratio may overinvest at substations with moderate generation-to-load ratios while failing to adequately "
    "represent reverse-power-flow constraints at substations with high distributed-PV penetration [2], [3]. Existing "
    "planning commonly relies on fixed or empirical CLR rules, whereas evolutionary multiobjective optimization can "
    "directly search the trade-off between annualized cost and supply adequacy [2]–[4]. Related evolutionary transfer "
    "studies also indicate that prior search knowledge can improve optimization efficiency [5]. However, random "
    "initialization and variation under strongly coupled engineering constraints can generate a large proportion of "
    "infeasible individuals.")

# ============ Introduction para 2 (§7.4) ============
setrun0(P[11],
    "This study proposes EPG-NSGA-II, which retains the non-dominated sorting and crowding-distance selection "
    "mechanisms of NSGA-II and introduces two engineering-prior operators: feasibility-aware warm-start sampling and "
    "a feasibility-restoring repair operator. The contributions are threefold. First, a bi-objective planning model "
    "jointly optimizes differentiated CLRs and storage power ratings under a load-weighted aggregate CLR band and "
    "station-level and upstream reverse-power-flow constraints, with EENS quantifying supply inadequacy under N-1 "
    "contingencies. Second, EPG-NSGA-II embeds engineering knowledge into population initialization and post-variation "
    "repair. Third, a public-data-informed synthetic county test system is used for comparative and ablation studies.")

# ============ Planning Model text: targeted edits (preserve inline math runs) ============
# P14: define K (spec §8.1); drop 'a set of'
repl(P[14], "The model considers a set of 110 kV substations.",
     "The model considers K 110-kV substations within a county-level grid, where K is the number of substations.")
repl(P[14], "110 kV substations", "110-kV substations")
# P22: storage not firm capacity (spec §8.6 wording) + county-level
repl(P[22], "Storage only absorbs reverse power flow, not credited as firm N-1 backup.",
     "Storage is modeled only as a reverse-power-flow mitigation resource and is not credited as firm capacity under "
     "N-1 contingencies.")

# ============ B. EPG-NSGA-II Procedure: formalize WarmStart/Repair (§9), Scheme B (§9.4) ============
# P24 procedure intro (§9.1): NSGA-II backbone [4]; define feasible aggregate-CLR set R (as inline text; a compact
# display equation for R is added in stage 2)
setrun0(P[24],
    "EPG-NSGA-II retains the NSGA-II backbone [4] and applies a feasibility-aware warm start before the first "
    "objective evaluation and an engineering repair operator after each variation step. Let R = (R₁,…,R_K) "
    "and P = (P₁,…,P_K). The feasible aggregate-CLR set ℛ collects all R that satisfy the box bounds "
    "R_j ∈ [1.2, 3.0] and the aggregate-CLR band (2).")
# P25 WarmStartSampling (§9.2) — unicode inline math
setrun0(P[25],
    "WarmStartSampling first draws R_j⁰ ∼ U(1.2, 3.0) and maps the sampled vector onto ℛ by rescaling "
    "it into the aggregate-CLR band. Given R, each storage rating is initialized as P_j ← min{6, [r_j⁰ "
    "− βR_jL_jᵖ]₊} MW, where [x]₊ = max{0, x}. If the upstream constraint (3) remains "
    "violated, additional storage is allocated to substations in descending order of residual reverse injection "
    "until (3) is satisfied or all storage ratings reach 6 MW.")
# P26 RepairProjection (§9.3 + Scheme B §9.4)
setrun0(P[26],
    "After variation, the repair operator clips the decision variables to their bounds and projects the CLR vector "
    "onto ℛ by the same rescaling. It then updates each storage rating as P_j ← min{6, max{P_j, [r_j⁰ "
    "− βR_jL_jᵖ]₊}} MW, followed by the same residual-based allocation for the upstream "
    "constraint, and performs a final feasibility test. If a constraint remains violated because all adjustable "
    "variables have reached their bounds, the residual violation is retained and handled by constraint-domination "
    "based on the total violation.")

# ============ Setup (§11.1), final numbering; K definition; mean-of-3-runs caliber ============
setrun0(P[29],
    "The benchmark is a public-data-informed synthetic county test system constructed from scaled IEEE 33-bus "
    "feeders [6] and PVGIS irradiance data [7]. Four system sizes are considered, K ∈ {10, 20, 30, 40}, where K "
    "denotes the number of 110-kV substations. The population size is N = 100, and each setting is evaluated over "
    "three independent runs. Standard NSGA-II and NSGA-III [8] are run for 800 generations as longer-budget "
    "references, whereas EPG-NSGA-II and its ablations are run for 400 generations.")

# ============ Table I caption + analysis (§11.2/§11.3) ============
setrun0(P[31], "TABLE I. NORMALIZED FINAL HV (MEAN OF THREE RUNS)")
setrun0(P[32],
    "Table I reports the mean final HV over three independent runs, normalized by the corresponding mean HV of "
    "standard NSGA-II for each K. EPG-NSGA-II matches or exceeds standard NSGA-II for all four system sizes, with "
    "the largest improvement of 4.8% at K=40. Repair-only closely tracks the full method, indicating that the repair "
    "operator accounts for most of the final-HV improvement in this experiment. Warm-start-only remains close to the "
    "baseline and is primarily associated with early-stage acceleration. NSGA-III yields lower final HV than standard "
    "NSGA-II for the tested bi-objective instances.")
# P33: 84-generation sentence (§11.4)
setrun0(P[33],
    "At K=40, EPG-NSGA-II attains the 800-generation standard-NSGA-II final HV after about 84 generations — roughly "
    "10.5% as many objective evaluations at the same population size.")

# ============ Application (§12.2) ============
setrun0(P[35],
    "The K=20 case is used as a representative county-level application. Its knee solution has a load-weighted "
    "aggregate CLR of 1.80, an annualized cost of 188.82 million CNY/yr, and an EENS of 114 MWh/yr. In the selected "
    "solution, storage is preferentially allocated to substations with higher generation-to-load ratios, whereas "
    "substations assigned lower CLRs rely more on fixed tie-line support (Fig. 1). Relative to the uniform R=2.0 "
    "baseline, the differentiated design reduces annualized cost by 13.1% at the same EENS level.")

# ============ Fig. 1 caption (§12.4) ============
setrun0(P[37],
    "Fig. 1.  K=20 differentiated CLR and storage allocation.")

# ============ Conclusion (§13) — drop feasibility-preserving; final HV ============
setrun0(P[39],
    "EPG-NSGA-II combines feasibility-aware warm-start sampling with an engineering repair operator for "
    "differentiated CLR optimization in a county-level 110-kV grid under high distributed PV penetration. In the "
    "synthetic test cases, the method reduces annualized cost at the same EENS level and achieves equal or higher "
    "final HV than standard NSGA-II, with the largest improvement of 4.8% at K=40. Future work will focus on "
    "calibration using real-grid data, energy-limited storage modeling, and more extensive statistical testing.")

# ============ non-dominated (body/algorithm only, NOT reference titles) ============
ref_idx = set()
for i, p in enumerate(d.paragraphs):
    if re.match(r'^\s*\[\d+\]', p.text):
        ref_idx.add(i)
for i, p in enumerate(d.paragraphs):
    if i in ref_idx:
        continue
    for r in p.runs:
        if "nondominated" in r.text:
            r.text = r.text.replace("nondominated", "non-dominated")

# ============ References: delete Nicolini (current [4]); relabel [5..9] -> [4..8]; complete [1] ============
allp = list(d.paragraphs)
refs = [p for p in allp if re.match(r'^\s*\[\d+\]', p.text)]
# delete Nicolini
for p in refs:
    if "Nicolini" in p.text:
        p._p.getparent().remove(p._p)
# relabel remaining refs sequentially by current number order (they are already in appearance order)
refs2 = [p for p in list(d.paragraphs) if re.match(r'^\s*\[\d+\]', p.text)]
refs2.sort(key=lambda p: int(re.match(r'^\s*\[(\d+)\]', p.text).group(1)))
for newnum, p in enumerate(refs2, 1):
    p.runs[0].text = re.sub(r'^\s*\[\d+\]', '[%d]' % newnum, p.runs[0].text, count=1)
# complete [1] Hou (§14.4)
for p in refs2:
    if "Hou" in p.text and "duck curve" in p.text:
        # rebuild run structure: [1] pre + journal(italic) + post
        pre = ('[1]  Q. Hou, N. Zhang, E. Du, M. Miao, F. Peng, and C. Kang, “Probabilistic duck curve in high '
               'PV penetration power system: Concept, modeling, and empirical analysis in China,” ')
        post = ', vol. 242, pp. 205–215, 2019, doi: 10.1016/j.apenergy.2019.03.067.'
        p.runs[0].text = pre
        if len(p.runs) >= 2:
            p.runs[1].text = "Applied Energy"
            p.runs[1].font.italic = True
            for r in p.runs[2:]:
                r.text = ""
            # ensure a post run exists
            rr = p.add_run(post); rr.font.size = Pt(7.5); rr.font.name = "Times New Roman"
        break
# shorten verbose PVGIS entry (§14.6 format consistency)
for p in refs2:
    if "PVGIS" in p.text:
        m = re.match(r'^\s*(\[\d+\])', p.text)
        lbl = m.group(1) if m else "[7]"
        p.runs[0].text = (lbl + "  European Commission, JRC, “Photovoltaic Geographical Information System "
                          "(PVGIS),” 2026. [Online]. Available: https://re.jrc.ec.europa.eu/pvg_tools/en/. "
                          "[Accessed: Jul. 2026].")
        for r in p.runs[1:]:
            r.text = ""
        break

# ============ STAGE 2: rebuild eq(2) [add δ] and eq(4) f2 [drop T_j] as native OMML ============
# eq (2): add county load-coincidence factor δ in the denominator (matches code DIVERSITY_FACTOR;
#         same δ as eq (3); preserves the reported δ-adjusted knee 1.80). Slash form -> single line.
set_eq_para(d.tables[1].rows[0].cells[1].paragraphs[0],
            r"1.8\le \sum_{j} R_j L_j^{p}\,/\,\bigl(\delta\textstyle\sum_{j} L_j^{p}\bigr)\le 2.2")
# eq (4) f2 line: drop undefined T_j -> f2 = sum_j EENS_j (tie-line support enters via S_{j,s} in eq (5))
set_eq_para(d.tables[3].rows[0].cells[1].paragraphs[2],
            r"f_2=\sum_{j}\text{EENS}_j")

# ============ STAGE 2b: Algorithm 1 is fully rebuilt in STAGE 4b (needs make_inline_omath, defined in STAGE 4) ============

# ============ Table I ablation label: 'Warm-start-only' -> 'Seeding-only' (Step 9) ============
for _t in d.tables:
    for _row in _t.rows:
        for _c in _row.cells:
            if "Warm-start-only" in _c.text:
                _hit = False
                for _pp in _c.paragraphs:
                    for _r in _pp.runs:
                        if "Warm-start-only" in _r.text:
                            _r.text = _r.text.replace("Warm-start-only", "Seeding-only"); _hit = True
                if not _hit and _c.paragraphs and _c.paragraphs[0].runs:  # split across runs
                    _c.paragraphs[0].runs[0].text = "Seeding-only"
                    for _r in _c.paragraphs[0].runs[1:]:
                        _r.text = ""

# ============ STAGE 3: hold 2 pages per §15.3 (cut redundant prose, not font) ============
# δ: introduce at (2) [advisor], remove the re-definition at (3); consistent single δ (code DIVERSITY_FACTOR)
setrun0(P[17],
    "Here δ = 0.85 is the county load-coincidence factor. Residual reverse power flow is limited at the station and "
    "upstream levels,")
setrun0(P[18],
    "where β = 0.85 is the transformer reverse-loading limit, R₂₂₀ = 1.8 the upper-level CLR, and η = 1.0 the "
    "reverse-power coincidence factor.")
# P22: drop redundant 'subject to ...' (constraints already stated by the equations)
setrun0(P[22],
    "Storage is modeled only as a reverse-power-flow mitigation resource and is not credited as firm capacity under "
    "N-1 contingencies.")
# P24: fold R/P definition into the R-set sentence
setrun0(P[24],
    "EPG-NSGA-II retains the NSGA-II backbone [4] and applies a feasibility-aware warm start before the first "
    "objective evaluation and an engineering repair operator after each variation step. The feasible aggregate-CLR "
    "set ℛ collects all R = (R₁,…,R_K) satisfying R_j ∈ [1.2, 3.0] and the aggregate-CLR band (2).")
# P26: reference the warm-start formula instead of repeating it
setrun0(P[26],
    "After variation, the repair operator applies the same clipping, projection onto ℛ, and residual-based storage "
    "update as the warm start, and then performs a final feasibility test. If a constraint remains violated because "
    "all adjustable variables have reached their bounds, the residual violation is retained and handled by "
    "constraint-domination based on the total violation.")

# --- further prose cuts per §15.3 (redundant natural language / repeated constraint text) ---
setrun0(P[11],
    "EPG-NSGA-II retains the standard variation, non-dominated sorting, and crowding-distance selection of NSGA-II "
    "while adding two problem-specific operators: engineering-prior population seeding and feasibility repair. The "
    "bi-objective model optimizes station-specific CLRs and storage ratings under aggregate-CLR and reverse-flow "
    "constraints, with EENS measuring N-1 supply inadequacy. Comparative and ablation studies are conducted on a "
    "public-data-informed county-level test system.")
setrun0(P[24],
    "EPG-NSGA-II keeps the NSGA-II backbone [4], adding a feasibility-aware warm start before the first evaluation and "
    "a repair operator after each variation. The feasible aggregate-CLR set is ℛ = {R : R_j ∈ [1.2, 3.0], band (2) "
    "holds}.")
setrun0(P[26],
    "After variation, the repair operator applies the same clipping, ℛ-projection, and storage update, then tests "
    "feasibility; any residual violation, arising when all adjustable variables are at their bounds, is handled by "
    "constraint-domination on the total violation.")
setrun0(P[35],
    "The K=20 case is a representative county-level application: its knee has a load-weighted aggregate CLR of 1.80, "
    "an annualized cost of 188.82 million CNY/yr, and an EENS of 114 MWh/yr. Storage concentrates at substations with "
    "higher generation-to-load ratios (Fig. 1). Relative to "
    "the uniform R=2.0 baseline, the differentiated design cuts annualized cost by 13.1% at the same EENS.")
setrun0(P[14],
    "The model considers K 110-kV substations within a county-level grid, where K is the number of substations. Each "
    "is assigned two variables: a capacity-load ratio and a storage power rating for absorbing reverse power flow. "
    "Current peak loads are projected to the planning year at 5% annual growth over five years; the projected peak of "
    "station j is Lⱼᵖ, used consistently in the capacity, reverse-power-flow, and supply-adequacy terms.")
setrun0(P[20],
    "Here, c_R and c_P are annualized unit costs and Cⱼˡᵒˢˢ, Cⱼᵒᵐ the annualized loss and O&M costs; f₂ is the total "
    "EENS, quantifying supply inadequacy under N-1 contingencies,")
setrun0(P[29],
    "The benchmark is a public-data-informed synthetic county test system from scaled IEEE 33-bus feeders [6] and "
    "PVGIS irradiance [7], with K ∈ {10, 20, 30, 40} 110-kV substations, population N = 100, and three independent "
    "runs per setting. Standard NSGA-II and NSGA-III [8] run 800 generations as longer-budget references; EPG-NSGA-II "
    "and its ablations run 400.")
# model explanations (unicode inline math; tightened)
setrun0(P[16],
    "Here, Pⱼ is the storage power rating and r̃ⱼ(Pⱼ) the residual reverse power flow after storage; the "
    "load-weighted aggregate CLR is constrained by band (2),")
setrun0(P[21],
    "In (5), Ωⱼ is the outage-state set, pₛ and hₛ the probability and duration of state s, and Sⱼ,ₛ the available "
    "supply from transformer capacity and fixed tie-line support.")
setrun0(P[22],
    "Storage mitigates reverse power flow but is not credited as firm capacity under N-1 contingencies.")
# P19 (Step 6.5): drop 'N-1 supply risk'
setrun0(P[19],
    "The two conflicting objectives are annualized cost and expected energy not supplied under N-1 contingencies,")
# intro para 1 (Step 5): new-operator framing; KEEP [5] with neutral prior-knowledge wording (not transfer optimization)
setrun0(P[10],
    "High distributed-PV penetration can cause midday export and reverse power flow at 110-kV substations, altering "
    "supply adequacy under N-1 contingencies [1]. A uniform CLR may overinvest at substations with moderate "
    "generation-to-load ratios yet fail to reflect reverse-flow limits at substations with high PV penetration "
    "[2], [3]. Evolutionary multiobjective optimization directly searches the cost–adequacy trade-off [2]–[4], and "
    "prior search knowledge can further improve its efficiency [5]. Under coupled engineering constraints, however, "
    "random initialization and variation generate many infeasible solutions.")
# WarmStart prose: tighten
setrun0(P[25],
    "WarmStartSampling draws Rⱼ⁰ ∼ U(1.2, 3.0), rescales the vector onto ℛ, and sets each storage rating Pⱼ ← "
    "min{6, [rⱼ⁰ − βRⱼLⱼᵖ]₊} MW, where [x]₊ = max{0, x}; if the upstream constraint (3) is still violated, storage "
    "is added to the substations with largest residual reverse injection until (3) holds.")
# table analysis: tighten
setrun0(P[32],
    "Table I reports the mean final HV over three runs, normalized by the mean HV of standard NSGA-II for each K. "
    "EPG-NSGA-II matches or exceeds the baseline for all four sizes, the largest gain being 4.8% at K=40. Repair-only "
    "closely tracks the full method (the repair operator drives most of the final-HV gain here), whereas "
    "warm-start-only stays near the baseline and mainly accelerates early search.")
# tie-line support already appears in eq (5) note; drop the trailing repetition in application if any
# conclusion: tighten
setrun0(P[39],
    "EPG-NSGA-II combines engineering-prior population seeding and feasibility repair for differentiated CLR "
    "optimization. In the county-level test cases, it reduces cost at equal EENS and achieves up to 4.8% higher "
    "mean final HV than standard NSGA-II. Future work will address real-grid calibration and energy-limited "
    "storage.")

# strip orphaned inline OMML math from prose paragraphs whose text was rewritten in Unicode
# (the original inline equations survive setrun0 and would render as trailing junk in Word too)
_M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
for _i in (14, 16, 18, 19, 20, 21):
    _pp = P[_i]._p
    for _m in _pp.findall('{%s}oMath' % _M_NS):
        _pp.remove(_m)

# figure smaller; reference line spacing slightly tighter (line spacing, not font size)
import struct as _st
_fig = ROOT + "/实验/有EA/results/figures/fig1_matrix.png"
_sh = d.inline_shapes[0]
_rid = _sh._inline.graphic.graphicData.pic.blipFill.blip.embed
with open(_fig, "rb") as _f:
    _data = _f.read()
d.part.related_parts[_rid]._blob = _data
_pw, _ph = _st.unpack(">II", _data[16:24])
import os as _os
_W = float(_os.environ.get("FIGW", "5.2"))
_sh.width = int(_W / 2.54 * 914400)
_sh.height = int(_sh.width * _ph / _pw)
# reference block: keep body 10pt intact; refs at conventional 7.0pt (IEEE), tight spacing, no inter-entry gaps
_REFPT = float(_os.environ.get("REFPT", "7.0"))
for p in d.paragraphs:
    if p.runs and p.runs[0].font.size and abs(p.runs[0].font.size.pt - 7.5) < 0.1:
        p.paragraph_format.line_spacing = 0.70
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        for r in p.runs:
            if r.font.size and abs(r.font.size.pt - 7.5) < 0.1:
                r.font.size = Pt(_REFPT)

# renderer-consistent reclaim: eq(4) has two stacked lines at 30pt exact line spacing (Σ padding) -> 23pt;
# and the figure spacer paragraph P[36] carries extra after-space. Both honored identically by Word/LibreOffice.
_eq4 = d.tables[3].rows[0].cells[1]
for _p in _eq4.paragraphs:
    _ls = _p.paragraph_format.line_spacing
    if _ls is not None and not isinstance(_ls, float) and _ls > Pt(24):
        _p.paragraph_format.line_spacing = Pt(23)
P[36].paragraph_format.space_before = Pt(0)
P[36].paragraph_format.space_after = Pt(0)

# ============ STAGE 4: inline body symbols -> native inline OMML (reproduce P[15]; sz=20 = body 10pt) ============
# The display equations (3),(4),(5) use default-italic math; reproduce that style inline (no \mathrm/\text),
# so every prose symbol matches the equations instead of Unicode/underscore text.
_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
_omath_cache = {}


def make_inline_omath(latex, half_pts=20):
    key = (latex, half_pts)
    if key not in _omath_cache:
        subprocess.run(["pandoc", "-f", "markdown", "-t", "docx", "-o", TMP_EQ],
                       input="$%s$" % latex, text=True, check=True)
        tree = etree.fromstring(zipfile.ZipFile(TMP_EQ).read("word/document.xml"))
        node = copy.deepcopy(tree.find(".//{%s}oMath" % M))
        _scale_math(node, half_pts)
        _omath_cache[key] = node
    return copy.deepcopy(_omath_cache[key])


def _mk_text_run(text, rpr_src):
    r = etree.Element("{%s}r" % W)
    if rpr_src is not None:
        r.append(copy.deepcopy(rpr_src))
    t = etree.SubElement(r, "{%s}t" % W)
    t.set(_XML_SPACE, "preserve")
    t.text = text
    return r


def _run0_halfpts(p, default=20):
    """Read the paragraph's own font size (half-points) so inline OMML matches it (body 20, abstract 18, …)."""
    if not p.runs:
        return default
    rpr = p.runs[0]._r.find("{%s}rPr" % W)
    if rpr is None:
        return default
    sz = rpr.find("{%s}sz" % W)
    return int(sz.get("{%s}val" % W)) if sz is not None else default


def set_mixed(p, template):
    """Rebuild a prose paragraph as interleaved text runs + inline OMML (template: $latex$ = math).
    Inline OMML is scaled to the paragraph's own font size so it never over/under-sizes the line."""
    pel = p._p
    rpr_src = p.runs[0]._r.find("{%s}rPr" % W) if p.runs else None
    hp = _run0_halfpts(p)
    for child in list(pel):
        if etree.QName(child).localname == "pPr":
            continue
        pel.remove(child)
    for i, seg in enumerate(template.split("$")):
        if i % 2 == 0:
            if seg:
                pel.append(_mk_text_run(seg, rpr_src))
        else:
            pel.append(make_inline_omath(seg, hp))


MIXED = {
    # --- Abstract (Step 3): alter (not reduce) supply adequacy; drop 'synthetic'; new operator names;
    #     'selected Pareto solution'; fix 84-gen reference. K=40 as inline OMML (auto-scaled to 9pt). ---
    7:  r"Abstract—High penetration of distributed photovoltaic (PV) generation can cause reverse power flow in "
        r"county-level 110-kV grids and alter supply adequacy under N-1 contingencies. This study jointly optimizes "
        r"differentiated capacity-load ratios (CLRs) and storage ratings by minimizing annualized cost and expected "
        r"energy not supplied (EENS). Engineering-Prior-Guided NSGA-II (EPG-NSGA-II) introduces engineering-prior "
        r"population seeding and feasibility repair to improve constrained search. On a public-data-informed "
        r"county-level test system, a selected Pareto solution reduces annualized cost by 13.1% at the same EENS. "
        r"For $K=40$, EPG-NSGA-II achieves a 4.8% higher mean final hypervolume than standard NSGA-II and reaches "
        r"the latter's 800-generation final level in about 84 generations.",
    # --- Planning model (Step 6.1): CLR defined; K, j, R_j, P_j, L_j^p as OMML. ---
    14: r"The county-level grid contains $K$ 110-kV substations. At station $j$, $R_j$ is the CLR, defined as "
        r"planned transformer capacity divided by forecast peak load; $P_j$ is the storage power rating; and "
        r"$L_j^{p}$ is the five-year forecast peak load obtained using 5% annual growth. CLRs vary across "
        r"substations but satisfy the load-weighted aggregate band.",
    16: r"Here $\tilde{r}_j(P_j)$ is the residual reverse power flow after storage, and the load-weighted aggregate "
        r"CLR is constrained by band (2),",
    17: r"Here $\delta = 0.85$ is the county load-coincidence factor. Residual reverse power flow is limited at the "
        r"station and upstream levels,",
    18: r"where $\beta = 0.85$ is the transformer reverse-loading limit, $R_{220} = 1.8$ the upper-level CLR, and "
        r"$\eta = 1.0$ the reverse-power coincidence factor.",
    # --- eq(4)/(5) parameter explanations, Step 6.6 (eq(5) sits between P20 and P21, so kept as two paras). ---
    20: r"Here, $c_R$ and $c_P$ are annualized unit costs, and $C_j^{loss}$ and $C_j^{om}$ are annual loss and "
        r"O&M costs; $f_2$ is the total EENS under N-1 contingencies,",
    21: r"In (5), $\Omega_j$ is the outage-state set, $p_s$ and $h_s$ the probability and duration of state $s$, "
        r"and $S_{j,s}$ the available transformer and tie-line supply.",
    # --- II-B intro + basic definitions (Step 8.2/8.3). ---
    24: r"EPG-NSGA-II retains the standard variation, non-dominated sorting, and crowding-distance selection of "
        r"NSGA-II [4]. It introduces engineering-prior population seeding for initialization and feasibility repair "
        r"for both the initial and offspring populations. Let $\operatorname{clip}(x,a,b)=\min\{b,\max\{a,x\}\}$ "
        r"and $[x]_+=\max\{0,x\}$.",
    # --- AggregateCLRAdjust note (Step 8.7): code-faithful = multiplicative rescaling (+ additive top-up), NOT projection. ---
    25: r"AggregateCLRAdjust rescales the CLR vector multiplicatively so that the load-weighted aggregate CLR "
        r"enters band (2), adding a top-up on the stations with headroom when the lower bound is still unmet, and "
        r"clips the result to $[R_{\min}, R_{\max}]$.",
    # --- FeasibilityRepair note (Step 8.8). ---
    26: r"FeasibilityRepair clips the variables, applies AggregateCLRAdjust, and raises storage — then $R_j$ — to "
        r"meet the per-station reverse-flow gate; residual upstream violations are repaired by allocating storage "
        r"to substations in descending order of residual reverse injection, and any infeasibility at the variable "
        r"bounds is passed to the total constraint violation.",
    # --- Setup (Step 9.1): keep [6] IEEE-33, [7] PVGIS, [8] NSGA-III (no deletions). ---
    29: r"The public-data-informed county-level test system is built from scaled IEEE 33-bus feeder data [6] and "
        r"PVGIS irradiance [7], with $K \in \{10, 20, 30, 40\}$ 110-kV substations and population $N = 100$, each "
        r"setting run three times. Standard NSGA-II and NSGA-III [8] run 800 generations as longer-budget "
        r"references; EPG-NSGA-II and its ablations run 400.",
    # --- Results (Step 9.3): Warm-start-only -> Seeding-only; keep NSGA-III. ---
    32: r"Table I normalizes the mean final HV by standard NSGA-II for each $K$. EPG-NSGA-II matches or exceeds the "
        r"baseline for all sizes, reaching $+4.8\%$ at $K=40$, while NSGA-III stays at or below standard NSGA-II. "
        r"Repair-only nearly matches the full method, whereas Seeding-only mainly accelerates early search.",
    33: r"At $K=40$, the mean EPG-NSGA-II curve reaches the 800-generation standard-NSGA-II final HV at about "
        r"generation 84 — roughly one-tenth of the evaluations.",
    # --- K=20 application (Step 10): 'selected Pareto solution'; keep 188.82 (real cost). ---
    35: r"For $K=20$, a selected Pareto solution has a load-weighted aggregate CLR of 1.80, an annualized cost of "
        r"188.82 million CNY/yr, and an EENS of 114 MWh/yr. Storage concentrates at substations with higher "
        r"generation-to-load ratios (Fig. 1), reducing cost by 13.1% versus the uniform $R=2.0$ baseline at equal "
        r"EENS.",
    # --- Fig. 1 caption (Step 10): compressed; K=20 as OMML (auto-scaled to caption size). ---
    37: r"Fig. 1.  $K=20$ differentiated CLR and storage allocation.",
}
if _os.environ.get("NOMATH") != "1":
    for _idx, _tmpl in MIXED.items():
        set_mixed(P[_idx], _tmpl)
    # optional: clamp inline-math paragraph line height (exact) so the math box can't inflate the line
    _MLS = _os.environ.get("MATHLS")
    if _MLS:
        for _idx in MIXED:
            P[_idx].paragraph_format.line_spacing = Pt(float(_MLS))

# ============ STAGE 4b: Algorithm 1 rebuilt (Step 8.4/8.5) — renamed + seeding/repair pseudocode; math as OMML ============
_AHP = 15  # 7.5 pt = algorithm body size


def _algo_line(para, tokens):
    """Rebuild one pseudocode line: 'kw'=bold keyword run, 't'=normal text run, 'm'=inline OMML (7.5pt)."""
    pel = para._p
    for _ch in list(pel):
        if etree.QName(_ch).localname == "pPr":
            continue
        pel.remove(_ch)
    for _kind, _val in tokens:
        if _kind == "m":
            pel.append(make_inline_omath(_val, _AHP))
        else:
            _r = etree.SubElement(pel, "{%s}r" % W)
            _rpr = etree.SubElement(_r, "{%s}rPr" % W)
            _rf = etree.SubElement(_rpr, "{%s}rFonts" % W)
            _rf.set("{%s}ascii" % W, "Times New Roman"); _rf.set("{%s}hAnsi" % W, "Times New Roman")
            if _kind == "kw":
                etree.SubElement(_rpr, "{%s}b" % W)
            for _tag in ("sz", "szCs"):
                _e = etree.SubElement(_rpr, "{%s}%s" % (W, _tag)); _e.set("{%s}val" % W, str(_AHP))
            _t = etree.SubElement(_r, "{%s}t" % W); _t.set(_XML_SPACE, "preserve"); _t.text = _val


_ALINES = [
    [("kw", "Input:"), ("t", " planning data "), ("m", "D"), ("t", ", population size "), ("m", "N"), ("t", ", and variable bounds.")],
    [("kw", "Output:"), ("t", " repaired initial population "), ("m", r"\mathcal{P}_0"), ("t", ".")],
    [("t", "1: "), ("m", r"\mathcal{P}_0 \leftarrow \varnothing")],
    [("t", "2: "), ("kw", "for "), ("m", r"i = 1, \ldots, N"), ("kw", " do")],
    [("t", "3:  "), ("m", r"\widetilde{\mathbf{R}}_i \sim \mathcal{U}([R_{\min}, R_{\max}])")],
    [("t", "4:  "), ("m", r"\mathbf{R}_i \leftarrow"), ("t", " AggregateCLRAdjust("), ("m", r"\widetilde{\mathbf{R}}_i"), ("t", ")")],
    [("t", "5:  "), ("m", r"\mathbf{P}_i \leftarrow \mathbf{0}"), ("t", "  (storage sized in FeasibilityRepair)")],
    [("t", "6:  "), ("m", r"\mathbf{x}_i \leftarrow"), ("t", " FeasibilityRepair("), ("m", r"\mathbf{R}_i, \mathbf{P}_i"), ("t", ")")],
    [("t", "7:  "), ("m", r"\mathcal{P}_0 \leftarrow \mathcal{P}_0 \cup \{\mathbf{x}_i\}")],
    [("t", "8: "), ("kw", "end for")],
    [("t", "9: "), ("kw", "return "), ("m", r"\mathcal{P}_0")],
    [("t", "FeasibilityRepair("), ("m", r"\mathbf{R}, \mathbf{P}"), ("t", "):")],
    [("t", "10:  "), ("m", r"\mathbf{R} \leftarrow"), ("t", " AggregateCLRAdjust("), ("m", r"\operatorname{clip}(\mathbf{R}, R_{\min}, R_{\max})"), ("t", ")")],
    [("t", "11:  "), ("m", r"P_j \leftarrow [r_j^0 - \beta R_j L_j^{p}]_+"), ("t", "  (capped at "), ("m", r"P_{\max}"), ("t", ")")],
    [("t", "12:  "), ("t", "Repair the upstream reverse-flow constraint in descending residual-injection order")],
    [("t", "13:  "), ("kw", "return "), ("m", r"(\mathbf{R}, \mathbf{P})")],
]
if _os.environ.get("NOMATH") != "1":
    _algo_cell = d.tables[5].rows[1].cells[0]
    _tmpl_ppr = _algo_cell.paragraphs[2]._p.find("{%s}pPr" % W)
    while len(_algo_cell.paragraphs) < len(_ALINES):
        _np = _algo_cell.add_paragraph()
        for _e in _np._p.findall("{%s}pPr" % W):
            _np._p.remove(_e)
        if _tmpl_ppr is not None:
            _np._p.insert(0, copy.deepcopy(_tmpl_ppr))
    _aparas = _algo_cell.paragraphs
    for _i, _toks in enumerate(_ALINES):
        _algo_line(_aparas[_i], _toks)
    for _extra in _algo_cell.paragraphs[len(_ALINES):]:
        _extra._p.getparent().remove(_extra._p)
    _cap = d.tables[5].rows[0].cells[0].paragraphs[0]
    if _cap.runs:
        _cap.runs[0].text = "Algorithm 1  Engineering-Prior Seeding and Feasibility Repair"
        for _r in _cap.runs[1:]:
            _r.text = ""

# ============ Global safety net: purge 'synthetic county' and 'N-1 supply risk' from body (not refs) ============
_SUB = [("synthetic county test system", "county-level test system"),
        ("synthetic county", "county-level"),
        ("N-1 supply risk", "supply inadequacy under N-1 contingencies")]
for _p in d.paragraphs:
    if re.match(r'^\s*\[\d+\]', _p.text):
        continue
    for _r in _p.runs:
        for _a, _b in _SUB:
            if _a in _r.text:
                _r.text = _r.text.replace(_a, _b)

# ============ STAGE 5: IEEE 3-per-row author grid (6 placeholder authors), matching the 范文 layout ============
from docx.oxml import OxmlElement as _OxE
from docx.oxml.ns import qn as _qn
from docx.enum.table import WD_TABLE_ALIGNMENT as _WTA
from docx.enum.text import WD_ALIGN_PARAGRAPH as _WAP

_ASZ = Pt(10)
_ORD = {1: "st", 2: "nd", 3: "rd", 4: "th", 5: "th", 6: "th"}
_AFF = ("School of Electrical Engineering", "China University of Mining and Technology", "Xuzhou, China")
# author 1 = real; 2-6 = clearly-fillable placeholders (same institution assumed; edit as needed)
_AUTHORS = [
    ("Heng Xu", _AFF[0], _AFF[1], _AFF[2], "ts24230188p31@cumt.edu.cn"),
    ("Second Author", _AFF[0], _AFF[1], _AFF[2], "coauthor2@cumt.edu.cn"),
    ("Third Author", _AFF[0], _AFF[1], _AFF[2], "coauthor3@cumt.edu.cn"),
    ("Fourth Author", _AFF[0], _AFF[1], _AFF[2], "coauthor4@cumt.edu.cn"),
    ("Fifth Author", _AFF[0], _AFF[1], _AFF[2], "coauthor5@cumt.edu.cn"),
    ("Sixth Author", _AFF[0], _AFF[1], _AFF[2], "coauthor6@cumt.edu.cn"),
]


def _style_run(r, italic=False):
    r.font.size = _ASZ
    r.font.name = "Times New Roman"
    r.font.italic = italic


def _fill_author(cell, idx, data):
    name, dept, univ, city, email = data
    p0 = cell.paragraphs[0]
    p0.alignment = _WAP.CENTER
    _style_run(p0.add_run(str(idx + 1)))
    _rs = p0.add_run(_ORD[idx + 1]); _style_run(_rs); _rs.font.superscript = True
    _style_run(p0.add_run(" " + name))
    for _text, _ital in [(dept, True), (univ, True), (city, False), (email, False)]:
        _p = cell.add_paragraph(); _p.alignment = _WAP.CENTER
        _style_run(_p.add_run(_text), _ital)
    for _p in cell.paragraphs:
        _p.paragraph_format.space_before = Pt(0)
        _p.paragraph_format.space_after = Pt(0)
        _p.paragraph_format.line_spacing = 1.0


_atbl = d.add_table(rows=2, cols=3)
_atbl.alignment = _WTA.CENTER
_atbl.autofit = False
_tblPr = _atbl._tbl.tblPr
_bd = _OxE("w:tblBorders")
for _e in ("top", "left", "bottom", "right", "insideH", "insideV"):
    _el = _OxE("w:" + _e); _el.set(_qn("w:val"), "nil"); _bd.append(_el)
_tblPr.append(_bd)
_lay = _OxE("w:tblLayout"); _lay.set(_qn("w:type"), "fixed"); _tblPr.append(_lay)
_sec0 = d.sections[0]
_cellw = int((_sec0.page_width - _sec0.left_margin - _sec0.right_margin) / 3)
for _k, _data in enumerate(_AUTHORS):
    _r, _c = divmod(_k, 3)
    _cell = _atbl.cell(_r, _c)
    _cell.width = _cellw
    _fill_author(_cell, _k, _data)
for _c in range(3):  # small gap above row 2 (matches 范文)
    _atbl.cell(1, _c).paragraphs[0].paragraph_format.space_before = Pt(8)
P[0]._p.addnext(_atbl._tbl)          # place grid right after the title (stays in the full-width section)
for _i in (5, 4, 3, 2, 1):           # drop the old single-affiliation author paragraphs
    _op = P[_i]._p
    _op.getparent().remove(_op)

d.save(OUT)
print("STAGE1..5 saved:", OUT)
# quick citation check
allp = list(d.paragraphs)
body = " ".join(p.text for p in allp if not re.match(r'^\s*\[\d+\]', p.text))
cites = sorted(set(int(x) for x in re.findall(r'\[(\d+)\]', body)))
reflabels = [int(re.match(r'^\s*\[(\d+)\]', p.text).group(1)) for p in allp if re.match(r'^\s*\[\d+\]', p.text)]
print("body cites:", cites, "| ref labels:", sorted(reflabels))
