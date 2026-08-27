#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Render the revised EPG-NSGA-II short paper to an IEEEtran-style Word .docx.
Revised per paper/初稿修正建议.md: II Problem Formulation (storage enters the
reverse-power constraints via the residual r_tilde; unified planning-year load L*;
tie-line T_j in EENS) + III EPG-NSGA-II (single Algorithm 1: main loop plus the
Repair procedure); down-toned, case-limited wording. Equations are native Word math
(OMML via pandoc); inline symbols are native runs. Output: epg-nsga-ii-paper_revised.docx.
Usage: python3 build_paper_docx.py"""
from __future__ import annotations
import os, subprocess, zipfile, copy
from lxml import etree
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL, WD_TAB_ALIGNMENT as TAB
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = "/home/roscy/ws_HengXU/徐州地区分布式新能源高渗透率地区110kV电网容载比弹性指标优化研究"
FIG = os.path.join(ROOT, "实验/有EA/results/figures/fig1_matrix.png")
OUT = os.path.join(ROOT, "paper/epg-nsga-ii-paper_revised_v2.docx")
TMP_EQ = "/tmp/claude-1000/-home-roscy-ws-HengXU-----------------110kV-------------/b58606b0-0593-4e06-b592-ccea2e68fc85/scratchpad/_eq.docx"
TNR, MONO = "Times New Roman", "Consolas"
GRAY = RGBColor(0x66, 0x66, 0x66)
M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
BODY = 10.0

# ---------- OMML (real Word equations via pandoc) ----------
def latex_omml(latex, display=False):
    md = ("$$%s$$" % latex) if display else ("$%s$" % latex)
    subprocess.run(["pandoc", "-f", "markdown", "-t", "docx", "-o", TMP_EQ],
                   input=md, text=True, check=True)
    xml = zipfile.ZipFile(TMP_EQ).read("word/document.xml")
    tree = etree.fromstring(xml)
    tag = "oMathPara" if display else "oMath"
    node = tree.find(".//{%s}%s" % (M, tag))
    node = copy.deepcopy(node)
    _scale_math(node, half_pts=20)   # 10 pt, matching the 10 pt body and the exemplar (Word honors this)
    return node


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _scale_math(node, half_pts=19):
    for mr in node.iter("{%s}r" % M):
        rpr = etree.SubElement(mr, "{%s}rPr" % W)
        for t in ("sz", "szCs"):
            e = etree.SubElement(rpr, "{%s}%s" % (W, t))
            e.set("{%s}val" % W, str(half_pts))
        mpr = mr.find("{%s}rPr" % M)
        mr.remove(rpr)
        mr.insert(1 if mpr is not None else 0, rpr)


def font(run, size, bold=False, ital=False, sc=False, mono=False, color=None):
    run.font.size = Pt(size); run.font.bold = bold; run.font.italic = ital
    run.font.small_caps = sc
    name = MONO if mono else TNR
    run.font.name = name
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia'):
        rf.set(qn(a), name)


def runs_md(p, text, size, **kw):
    for i, seg in enumerate(text.split("**")):
        if not seg:
            continue
        font(p.add_run(seg), size, bold=(i % 2 == 1) or kw.get("bold", False),
             ital=kw.get("ital", False), color=kw.get("color"))


def P(doc, text="", size=BODY, bold=False, ital=False, align=AL.JUSTIFY, before=0, after=0,
      line=1.0, indent=False, color=None, md=True):
    p = doc.add_paragraph(); p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = line
    if indent:
        p._p.get_or_add_pPr().get_or_add_ind().set(qn('w:firstLineChars'), '150')
    if text:
        if md and not ital:
            runs_md(p, text, size, bold=bold, color=color)
        else:
            font(p.add_run(text), size, bold=bold, ital=ital, color=color)
    return p


def rich(p, tokens, size=BODY):
    """typed runs: t=plain, kw=bold, vr=italic var, sb=subscript, sp=superscript, cm=gray italic."""
    for kind, txt in tokens:
        r = p.add_run(txt)
        if kind == "kw":
            font(r, size, bold=True)
        elif kind == "vr":
            font(r, size, ital=True)
        elif kind == "sb":
            font(r, size); r.font.subscript = True
        elif kind == "sp":
            font(r, size); r.font.superscript = True
        elif kind == "cm":
            font(r, size, ital=True, color=GRAY)
        else:
            font(r, size)


def PT(doc, tokens, indent=False, after=2):
    p = doc.add_paragraph(); p.alignment = AL.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(after); pf.line_spacing = 1.0
    if indent:
        p._p.get_or_add_pPr().get_or_add_ind().set(qn('w:firstLineChars'), '150')
    rich(p, tokens)
    return p


def spacer(doc):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.0
    font(p.add_run(" "), 2)
    return p


def _vcenter(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement('w:vAlign'); va.set(qn('w:val'), 'center'); tcPr.append(va)


def eq(doc, latex, num=None):
    """Numbered DISPLAY equation: borderless 3-col table [pad | oMathPara | (n)].
    latex may be a str (single line) or a list of str (stacked lines sharing one
    number) -- each line is its own single-line oMathPara, so the merged block reuses
    the already-Word-verified display primitive (no eqArr)."""
    lines = latex if isinstance(latex, (list, tuple)) else [latex]
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cells = t.rows[0].cells
    for i, ltx in enumerate(lines):
        pm = cells[1].paragraphs[0] if i == 0 else cells[1].add_paragraph()
        pm.alignment = AL.CENTER
        pm.paragraph_format.space_before = Pt(0); pm.paragraph_format.space_after = Pt(0)
        pm._p.append(latex_omml(ltx, display=True))
    pn = cells[2].paragraphs[0]; pn.alignment = AL.RIGHT
    if num:
        font(pn.add_run("(" + num + ")"), BODY)
    for c in cells:
        _vcenter(c)
    dxa = int(8.4 / 2.54 * 1440)
    t.autofit = False; t.allow_autofit = False
    tblPr = t._tbl.tblPr
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed'); tblPr.append(layout)
    w = OxmlElement('w:tblW'); w.set(qn('w:type'), 'dxa'); w.set(qn('w:w'), str(dxa)); tblPr.append(w)
    widths = (0.2, 7.2, 1.0)
    for c, cm_w in zip(cells, widths):
        c.width = Cm(cm_w)
    grid = t._tbl.find(qn('w:tblGrid'))
    for gc, cm_w in zip(grid.findall(qn('w:gridCol')), widths):
        gc.set(qn('w:w'), str(int(cm_w / 2.54 * 1440)))
    return t


def set_page(sec):
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.left_margin = sec.right_margin = Inches(0.62)
    sec.top_margin = Inches(0.75); sec.bottom_margin = Inches(1.0)


def set_cols(sec, num, space=360):
    sectPr = sec._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols'); sectPr.append(cols)
    cols.set(qn('w:num'), str(num)); cols.set(qn('w:space'), str(space)); cols.set(qn('w:equalWidth'), '1')


def new_section(doc, num):
    s = doc.add_section(WD_SECTION.CONTINUOUS); set_page(s)
    set_cols(s, num, 360 if num == 2 else 720)
    return s


def H1(doc, label):
    p = doc.add_paragraph(); p.alignment = AL.CENTER
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(1)
    font(p.add_run(label), 9.0, sc=True, bold=False)


def H2(doc, label):
    p = doc.add_paragraph(); p.alignment = AL.LEFT
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(1)
    font(p.add_run(label), 9.0, ital=True)


def caption(doc, text, before=3, keep_next=False):
    p = doc.add_paragraph(); p.alignment = AL.JUSTIFY
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = keep_next
    font(p.add_run(text), 8.0)


def fix_table_width(t, cm):
    dxa = int(cm / 2.54 * 1440)
    t.autofit = False; t.allow_autofit = False
    tblPr = t._tbl.tblPr
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed'); tblPr.append(layout)
    w = OxmlElement('w:tblW'); w.set(qn('w:type'), 'dxa'); w.set(qn('w:w'), str(dxa)); tblPr.append(w)
    for row in t.rows:
        for c in row.cells:
            c.width = Cm(cm / len(row.cells))


def _border(tc, edges, sz=8):
    tcPr = tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for e in edges:
        x = OxmlElement('w:' + e)
        x.set(qn('w:val'), 'single'); x.set(qn('w:sz'), str(sz)); x.set(qn('w:space'), '0'); x.set(qn('w:color'), '000000')
        b.append(x)
    tcPr.append(b)


def _set_col_widths(t, widths_cm):
    t.autofit = False; t.allow_autofit = False
    tblPr = t._tbl.tblPr
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed'); tblPr.append(layout)
    w = OxmlElement('w:tblW'); w.set(qn('w:type'), 'dxa')
    w.set(qn('w:w'), str(int(sum(widths_cm) / 2.54 * 1440))); tblPr.append(w)
    grid = t._tbl.find(qn('w:tblGrid'))
    for gc, cm in zip(grid.findall(qn('w:gridCol')), widths_cm):
        gc.set(qn('w:w'), str(int(cm / 2.54 * 1440)))
    for row in t.rows:
        for c, cm in zip(row.cells, widths_cm):
            c.width = Cm(cm)


def _tight_cell(c):
    pf = c.paragraphs[0].paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.0
    c.paragraphs[0].alignment = AL.CENTER


def three_line_table(doc, header, rows, size=8.0, width_cm=None, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = (width_cm is None and col_widths is None)
    for j, h in enumerate(header):
        c = t.rows[0].cells[j]; _tight_cell(c)
        font(c.paragraphs[0].add_run(h), size, bold=True)
        _border(c._tc, ["top", "bottom"], sz=10)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[i + 1].cells[j]; _tight_cell(c)
            font(c.paragraphs[0].add_run(str(v)), size)
            if i == len(rows) - 1:
                _border(c._tc, ["bottom"], sz=10)
    # keep the whole (small) table together so it never splits across a column
    for ridx, row in enumerate(t.rows):
        row._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))
        if ridx < len(t.rows) - 1:
            for c in row.cells:
                c.paragraphs[0].paragraph_format.keep_with_next = True
    if col_widths is not None:
        _set_col_widths(t, col_widths)
    elif width_cm is not None:
        fix_table_width(t, width_cm)
    return t


def _edges(tc, spec):
    tcPr = tc.get_or_add_tcPr(); b = OxmlElement('w:tcBorders')
    for e, sz in spec.items():
        x = OxmlElement('w:' + e)
        x.set(qn('w:val'), 'single'); x.set(qn('w:sz'), str(sz))
        x.set(qn('w:space'), '0'); x.set(qn('w:color'), '000000')
        b.append(x)
    tcPr.append(b)


def _algo_lines(cell, lines, head_tokens=None):
    """Render hanging-indent algorithm lines into a cell; optional bold header line."""
    seq = ([head_tokens] if head_tokens else []) + list(lines)
    first = True
    for ln in seq:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        is_head = first and head_tokens is not None
        first = False
        pf = p.paragraph_format
        pf.space_before = Pt(1 if is_head else 0); pf.space_after = Pt(0); pf.line_spacing = 1.0
        ind = p._p.get_or_add_pPr().get_or_add_ind()
        ind.set(qn('w:left'), '300'); ind.set(qn('w:hanging'), '300')
        rich(p, ln, size=7.8)


def algo_box(doc, title, lines, proc_head=None, proc_lines=None):
    """LaTeX-algorithmic look, native Word: three-line-ruled 1-col table.
    With proc_head/proc_lines, a Repair procedure is placed in its own row under an
    internal divider, so the table may break only at the procedure boundary (avoids a
    tall cantSplit block jumping a whole column)."""
    merged = proc_lines is not None
    t = doc.add_table(rows=3 if merged else 2, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    head = t.rows[0].cells[0]; body = t.rows[1].cells[0]
    _edges(head._tc, {"top": 12, "bottom": 4})
    if merged:
        proc = t.rows[2].cells[0]
        _edges(proc._tc, {"top": 4, "bottom": 12})   # internal divider + bottom rule
    else:
        _edges(body._tc, {"bottom": 12})
    for row in t.rows:
        row._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))
    head.paragraphs[0].paragraph_format.keep_with_next = True
    p0 = head.paragraphs[0]
    p0.paragraph_format.space_before = Pt(1); p0.paragraph_format.space_after = Pt(1)
    if ":" in title:
        no, name = title.split(":", 1)
        font(p0.add_run(no.strip() + " "), 8.3, bold=True)
        font(p0.add_run(name.strip()), 8.3)
    else:
        font(p0.add_run(title), 8.3, bold=True)
    _algo_lines(body, lines)
    if merged:
        _algo_lines(proc, proc_lines, head_tokens=proc_head)
    fix_table_width(t, 8.4)
    return t


def _grid_widths(t, widths_cm):
    t.autofit = False; t.allow_autofit = False
    tblPr = t._tbl.tblPr
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed'); tblPr.append(layout)
    w = OxmlElement('w:tblW'); w.set(qn('w:type'), 'dxa')
    w.set(qn('w:w'), str(int(sum(widths_cm) / 2.54 * 1440))); tblPr.append(w)
    grid = t._tbl.find(qn('w:tblGrid'))
    for gc, cm in zip(grid.findall(qn('w:gridCol')), widths_cm):
        gc.set(qn('w:w'), str(int(cm / 2.54 * 1440)))


def _cell_margins(t, lr=25, tb=6):
    tblPr = t._tbl.tblPr
    m = OxmlElement('w:tblCellMar')
    for side, val in (('top', tb), ('bottom', tb), ('left', lr), ('right', lr)):
        e = OxmlElement('w:' + side); e.set(qn('w:w'), str(val)); e.set(qn('w:type'), 'dxa'); m.append(e)
    tblPr.append(m)


def hv_matrix_table(doc, rows, size=8.0):
    """Two-level-header HV table: K x (gen 100 | converged) x (NSGA-II, Warm, Repair).
    Normalized to EPG-NSGA-II = 1.000 per K (reference stated in the caption)."""
    ncol = 7
    t = doc.add_table(rows=2 + len(rows), cols=ncol)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _grid_widths(t, [0.6, 1.2, 0.98, 1.02, 1.2, 0.98, 1.02])   # ~7.0 cm, fits one column
    _cell_margins(t, lr=22, tb=6)                              # tight margins so 0.xxx never wraps
    # group-header row (row 0): blank | "gen 100" (cols 1-3) | "converged" (cols 4-6)
    g1 = t.cell(0, 1).merge(t.cell(0, 2)).merge(t.cell(0, 3))
    _tight_cell(g1); font(g1.paragraphs[0].add_run("HV at gen. 100"), size, bold=True)
    g2 = t.cell(0, 4).merge(t.cell(0, 5)).merge(t.cell(0, 6))
    _tight_cell(g2); font(g2.paragraphs[0].add_run("HV at convergence"), size, bold=True)
    _tight_cell(t.cell(0, 0))
    subs = ["K", "NSGA-II", "Warm", "Repair", "NSGA-II", "Warm", "Repair"]
    for j, s in enumerate(subs):
        c = t.cell(1, j); _tight_cell(c); font(c.paragraphs[0].add_run(s), size, bold=True)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.cell(2 + i, j); _tight_cell(c); font(c.paragraphs[0].add_run(str(v)), size)
    # three-line rules: toprule, cmidrules under group labels, midrule under sub-headers, bottomrule
    for j in range(ncol):
        _border(t.cell(0, j)._tc, ["top"], sz=10)
    for j in range(1, ncol):
        _border(t.cell(0, j)._tc, ["bottom"], sz=6)
    for j in range(ncol):
        _border(t.cell(1, j)._tc, ["bottom"], sz=8)
    last = 2 + len(rows) - 1
    for j in range(ncol):
        _border(t.cell(last, j)._tc, ["bottom"], sz=10)
    # keep the whole table together so it does not split across the column break
    for ridx, r in enumerate(t.rows):
        r._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))
        if ridx < len(t.rows) - 1:
            for c in r.cells:
                c.paragraphs[0].paragraph_format.keep_with_next = True
    return t


# ================= content (revised per 初稿修正建议.md) =================
TITLE = ("Engineering-Prior-Guided NSGA-II for Differentiated Capacity-Load Ratio "
         "Optimization in County-Level 110-kV Grid Planning under High Distributed PV "
         "Penetration")
AUTHORS = "Heng Xu"
AFFIL_LINES = ["School of Electrical Engineering", "China University of Mining and Technology",
               "Xuzhou, Jiangsu, China", "ts24230188p31@cumt.edu.cn"]

ABSTRACT = ("High penetration of distributed photovoltaic (PV) generation turns county-level 110-kV substations into "
            "bidirectional hubs: a uniform capacity-load ratio (CLR) over-invests at some stations while "
            "under-constraining reverse power flow at others, and transformer capacity must still preserve supply "
            "adequacy under N-1 contingencies. This study formulates a multiobjective planning model that coordinates "
            "multiple 110-kV substations within a county and jointly optimizes differentiated capacity-load ratios and "
            "storage power ratings, minimizing annualized cost and expected energy not supplied (EENS) subject to a "
            "load-weighted aggregate CLR band and substation- and upstream-level reverse-power-flow constraints. The "
            "model is solved by Engineering-Prior-Guided NSGA-II (EPG-NSGA-II), which combines feasibility-preserving "
            "warm-start sampling with a constraint-ordered repair projection to improve constrained search. On the "
            "public-data-informed synthetic county test system, the differentiated design reduces annualized cost by "
            "13.1% relative to the uniform R=2.0 baseline at the same EENS level. Across county sizes K=10–40, "
            "EPG-NSGA-II's converged normalized hypervolume rises from near-parity at small K to +4.8% over standard "
            "NSGA-II at K=40, and it reaches the standard-NSGA-II converged front with about 9.5× fewer "
            "evaluations, the advantage widening with scale.")
INDEX_BODY = ("Capacity-load ratio, distributed photovoltaic generation, multiobjective optimization, reverse power "
              "flow, substation planning.")

INTRO = [
    "High penetration of distributed PV generation is turning county-level 110-kV grids into bidirectional networks: "
    "around midday local PV can exceed demand, so transformers experience reverse power flow and supply adequacy under "
    "N-1 contingencies shifts with net load [3]. A uniform capacity-load ratio (e.g., 2.0 across the whole county) "
    "over-invests where load and PV are modest and under-constrains reverse power flow at substations with high "
    "distributed-PV penetration [1], [2]; existing planning sets this ratio from fixed or empirical rules [1], [2], "
    "[4]. Multiobjective evolutionary algorithms instead search the cost-risk trade-off directly [5], and NSGA-II [6] "
    "variants are common for distribution planning and storage siting. Under strong constraints, however, random "
    "initialization and variation produce many infeasible individuals, so standard NSGA-II wastes much of its budget "
    "outside the feasible region, and the waste grows as the region tightens.",
    "This study proposes EPG-NSGA-II. Rather than changing nondominated sorting, it embeds engineering priors where "
    "infeasible individuals arise: a feasibility-preserving warm start that seeds the population inside the engineering "
    "bounds, and a constraint-ordered repair projection that maps offspring back toward feasibility after variation. "
    "Here, differentiated capacity-load ratio optimization assigns different ratios to substations within a county "
    "according to their load, distributed-PV, and network-support conditions, while keeping the load-weighted "
    "aggregate ratio within a prescribed county-level band; the capacity-load ratio is defined as the ratio of planned "
    "transformer capacity to forecast peak load. Contributions: 1) a compact multiobjective model with differentiated "
    "capacity-load ratios under an aggregate-CLR band and reverse-power-flow constraints and an EENS-based "
    "supply-adequacy objective under N-1 contingencies; 2) an engineering-prior-guided NSGA-II with "
    "feasibility-preserving warm-start sampling and a repair projection; and 3) a public-data-informed synthetic "
    "county test system with an ablation isolating the repair operator.",
]

# ---- II Problem Formulation connective lines (typed runs) ----
D_VARS = [("t", "where "), ("vr", "R"), ("sb", "j"), ("t", " is the planned capacity-load ratio of station "),
          ("vr", "j"), ("t", " and "), ("vr", "P"), ("sb", "j"),
          ("t", " is the installed storage power rating that absorbs local reverse power flow. The planning-year peak load is")]
AFTER_L = [("t", "with "), ("vr", "L"), ("sb", "j"), ("t", " the current peak, "), ("vr", "g"),
           ("t", " = 5% the annual load growth, and "), ("vr", "n"), ("t", " = 5 the horizon; the planning-year peak "),
           ("vr", "L"), ("sb", "j"), ("sp", "p"),
           ("t", " is used throughout, in the capacity, reverse-power, and risk terms. With "),
           ("vr", "r"), ("sb", "j"), ("sp", "0"),
           ("t", " the midday reverse-power injection before storage, the residual after storage is")]
AFTER_R = [("t", "The variables are bounded, "), ("vr", "R"), ("sb", "j"), ("t", " ∈ [1.2, 3.0] and "),
           ("vr", "P"), ("sb", "j"),
           ("t", " ∈ [0, 6] MW. The load-weighted aggregate capacity-load ratio is held within a planning band,")]
AFTER_AGG = [("t", "the ratio of total installed capacity to the county coincident peak, with δ = 0.85 the load "
                   "coincidence factor (stations do not peak simultaneously). The residual reverse power flow is limited "
                   "at each station and at the upstream interface,")]
AFTER_UP = [("t", "where β = 0.85 is the transformer reverse-loading limit, "), ("vr", "R"), ("sb", "220"),
            ("t", " = 1.8 the upper-level capacity-load ratio, and η = 1.0 the reverse coincidence factor.")]
OBJ_INTRO = "Two conflicting objectives are minimized, the annualized cost and the N-1 supply risk,"
AFTER_F1 = [("t", "where "), ("vr", "c"), ("sb", "R"), ("t", " and "), ("vr", "c"), ("sb", "P"),
            ("t", " are annualized unit costs of transformer capacity and storage power, and "), ("vr", "C"), ("sb", "j"),
            ("sp", "loss"), ("t", ", "), ("vr", "C"), ("sb", "j"), ("sp", "om"),
            ("t", " the annualized bidirectional-loss and O&M costs. The capacity term is an annualized capacity "
                  "proxy, not an incremental-investment accounting model. The N-1 supply risk "), ("vr", "f"), ("sb", "2"),
            ("t", " is the total expected energy not supplied (EENS),")]
AFTER_EENS = [("t", "where the single-element (N-1) outage states "), ("vr", "s"), ("t", " of station "), ("vr", "j"),
              ("t", " have probability "), ("vr", "p"), ("sb", "s"), ("t", " and duration "), ("vr", "h"), ("sb", "s"),
              ("t", ", and "), ("vr", "S"), ("sb", "j,s"),
              ("t", " is the available supply under outage "), ("vr", "s"),
              ("t", ", set by the transformer capacity and the exogenous tie-line support "),
              ("vr", "T"), ("sb", "j"),
              ("t", " (a fixed input, not optimized here). Storage is counted only as reverse-power-flow absorption and is "
                    "not credited as firm N-1 backup, so it enters the reverse-power-flow constraints (3) and (5) rather "
                    "than "), ("vr", "f"), ("sb", "2"),
              ("t", ". The optimization is subject to the variable bounds and (4), (5).")]

# ---- III method prose ----
III_A = ("EPG-NSGA-II keeps the NSGA-II backbone [6] (constraint-dominated nondominated sorting with crowding "
         "distance) and adds two engineering-prior operators (Algorithm 1). Both act only on the decision variables, "
         "not the objective, so they cost no evaluations: the warm start seeds a near-feasible initial population, and "
         "the repair projection is applied after every variation step.")
III_B = [("t", "The warm start (Algorithm 1, line 1) builds each individual from the engineering bounds: it samples "
                "each ratio "), ("vr", "R"), ("sb", "j"),
          ("t", " within [1.2, 3.0], projects the aggregate ratio into band (4), and allocates storage to stations in "
                "descending order of residual reverse-power-flow pressure "), ("vr", "r"), ("sb", "j"),
          ("sp", "0"), ("t", "/("), ("vr", "R"), ("sb", "j"), ("vr", "L"), ("sb", "j"), ("sp", "p"),
          ("t", "), seeding the population near the feasible region.")]
III_C = ("After variation, RepairProjection (Algorithm 1, lines 2 and 6) performs a constraint-ordered projection and "
         "a final feasibility check, without evaluating the objectives: it clips each ratio to [1.2, 3.0] and each "
         "storage rating to [0, 6] MW and rescales the ratios into band (4); at each station it raises storage first "
         "and lifts the ratio only when storage saturates, enforcing the residual reverse-power-flow gate (5); it "
         "repairs the upstream interface constraint (5); and it re-checks band (4). If the check still fails because "
         "all adjustable variables are at their bounds, the residual violation is passed to the constraint-domination "
         "ranking, so the projection improves feasibility preservation but does not guarantee it.")

ALGO1 = [
    [("kw", "Input: "), ("t", "station data "), ("vr", "D"), ("t", "; ratio and storage bounds; aggregate band; "
     "reverse-power params ("), ("vr", "R"), ("sb", "220"), ("t", ", β, δ, η); population "), ("vr", "N"),
     ("t", "; generations "), ("vr", "G"), ("t", ".")],
    [("kw", "Output: "), ("t", "feasible nondominated set "), ("vr", "A"), ("t", ".")],
    [("t", "1: "), ("vr", "P"), ("t", " ← WarmStartSampling("), ("vr", "D"), ("t", ", "), ("vr", "N"), ("t", ")")],
    [("t", "2: "), ("vr", "P"), ("t", " ← RepairProjection("), ("vr", "P"), ("t", ")")],
    [("t", "3: evaluate "), ("vr", "f"), ("sb", "1"), ("t", ", "), ("vr", "f"), ("sb", "2"),
     ("t", " for all individuals in "), ("vr", "P")],
    [("t", "4: "), ("kw", "for "), ("vr", "t"), ("t", " = 1 "), ("kw", "to"), ("t", " "), ("vr", "G"), ("t", " "), ("kw", "do")],
    [("t", "5:   "), ("vr", "Q"), ("t", " ← selection, crossover, mutation of "), ("vr", "P")],
    [("t", "6:   "), ("vr", "Q"), ("t", " ← RepairProjection("), ("vr", "Q"), ("t", ")   "),
     ("cm", "(no objective call)")],
    [("t", "7:   evaluate "), ("vr", "f"), ("sb", "1"), ("t", ", "), ("vr", "f"), ("sb", "2"),
     ("t", " for all individuals in "), ("vr", "Q")],
    [("t", "8:   "), ("vr", "U"), ("t", " ← "), ("vr", "P"), ("t", " ∪ "), ("vr", "Q")],
    [("t", "9:   "), ("vr", "P"), ("t", " ← nondominated sort + crowding select ("), ("vr", "U"), ("t", ", "),
     ("vr", "N"), ("t", ")")],
    [("t", "10: "), ("kw", "end for")],
    [("t", "11: "), ("vr", "A"), ("t", " ← feasible nondominated individuals in "), ("vr", "P")],
    [("t", "12: "), ("kw", "return"), ("t", " "), ("vr", "A")],
]

SETUP = ("The benchmark is a public-data-informed synthetic county test system from scaled IEEE 33-bus feeders [8] and PVGIS "
         "irradiance [9] for three sites, with fixed seeds. We use population 100, county sizes K ∈ {10, 20, 30, 40}, "
         "and three seeds per configuration, with a common hypervolume (HV) reference point within each K. Function "
         "evaluations scale with generations (fixed population); the two baselines, standard NSGA-II and NSGA-III [7], "
         "run to 800 generations and EPG-NSGA-II to 400.")

EFF = [
    "Table I reports converged hypervolume normalized to standard NSGA-II = 1.000 at each county size K. EPG-NSGA-II "
    "is at or above the baseline for every K, and the margin widens with scale, from +0.7% at K=10 to +4.8% at K=40. "
    "The repair-only variant tracks the full method almost exactly (+0.4% to +4.8%), so the repair projection "
    "accounts for most of the converged-quality gain; warm-start-only instead stays at or slightly below the "
    "baseline at convergence (-0.4% to -0.2%), since its benefit is early acceleration rather than final quality "
    "(below). NSGA-III stays below standard NSGA-II throughout (-0.2% to -2.1%), consistent with its reference-point "
    "machinery giving little benefit on this bi-objective problem; it is included only as a second evolutionary "
    "reference.",
    "The warm start instead accelerates early search. We measure this by the generation at which EPG-NSGA-II first "
    "reaches standard NSGA-II's converged (generation-800) hypervolume; with a fixed population, generations equal "
    "function evaluations up to a constant. At K=40 this takes 84 generations, about 9.5 times fewer evaluations. "
    "Warm-start-only alone already leads standard NSGA-II at generation 100 for K ≥ 20, by a margin that grows with "
    "scale (up to about 6% at K=30), then relaxes to baseline quality at convergence.",
]

TAB_CAP = ("TABLE I.  Converged hypervolume by algorithm and county size K, normalized to standard NSGA-II = 1.000 "
           "per K (higher is better). Warm-start-only and Repair-only are the warm-start and repair-projection "
           "ablations; NSGA-III is a second evolutionary reference.")
TAB_HEAD = ["Method", "K=10", "K=20", "K=30", "K=40"]
TAB_ROWS = [["NSGA-II", "1.000", "1.000", "1.000", "1.000"],
            ["NSGA-III", "0.998", "0.991", "0.979", "0.985"],
            ["Warm-start-only", "0.996", "0.991", "0.988", "0.998"],
            ["Repair-only", "1.004", "1.006", "1.008", "1.048"],
            ["EPG-NSGA-II", "1.007", "1.007", "1.009", "1.048"]]

APP = ("K=20 is used as a representative county size for the application study; the algorithm study spans K=10–40 to "
       "cover small-to-large county grids. At K=20, the knee of the returned Pareto front operates at a load-weighted "
       "aggregate capacity-load ratio of 1.80, with an annualized cost of about 188.82 million CNY/yr and EENS 114 "
       "MWh/yr. The selected configuration is spatially differentiated (Fig. 1): substations with high "
       "generation-to-load ratios take storage, low-ratio stations rely on tie-line support, and no station exceeds a "
       "ratio of 2.0 (maximum 1.96). A conventional uniform R=2.0 design corresponds to a load-weighted aggregate "
       "ratio of 2.35, above the planning band, reflecting the over-provisioning that differentiated planning removes; "
       "at the same EENS level the differentiated configuration reduces the annualized cost by 13.1% relative to that "
       "baseline in this synthetic case.")

FIG_CAP = ("Fig. 1.  Differentiated planning configuration on the synthetic county test system (K=20 knee). Substations "
           "are ordered by generation-to-load ratio; each cell gives the station capacity-load ratio R and storage "
           "(maximum R = 1.96).")
CONC = ("On the synthetic county test system, EPG-NSGA-II's feasibility-preserving warm start and repair projection "
        "reduced annualized cost against a uniform-ratio baseline at the same EENS level, and improved the converged "
        "hypervolume over standard NSGA-II by a margin that widens with county size, reaching +4.8% with about 9.5× "
        "fewer evaluations at K=40. Real-grid calibration, energy-limited storage, and statistical significance tests "
        "remain future work.")

REFS = [
    ('J. Xiao, R. Liu, and M. Long, "Quantitative analysis of the relationship between substation capacity-load ratio and the next-level medium-voltage distribution network," ', 'Electric Power Construction', ', vol. 36, no. 11, pp. 45-50, 2015 (in Chinese).'),
    ('K. Wang, Z. He, J. Ye, Y. Lin, and Y. Ma, "Capacity-load-ratio optimal configuration for coordinated planning of multi-voltage-level grids," ', 'Southern Power System Technology', ', vol. 19, no. 7, pp. 131-139, 2025 (in Chinese).'),
    ('Q. Hou, N. Zhang, E. Du, M. Miao, F. Peng, and C. Kang, "Probabilistic duck curve in high PV penetration power system," ', 'Applied Energy', ', 2019.'),
    ('T. Lin, G. Wu, S. Lai, H. Hu, and Z. Hu, "Calculation of distribution network PV hosting capacity considering source-load uncertainty and active management," ', 'Electronics', ', vol. 13, art. 4048, 2024.'),
    ('M. Nicolini, "Multi-objective genetic algorithms in designing redundant water distribution systems," in ', 'Proc. 2025 IEEE Int. Conf. MIND', ', 2025, doi: 10.1109/MIND67540.2025.11351874.'),
    ('K. Deb, A. Pratap, S. Agarwal, and T. Meyarivan, "A fast and elitist multiobjective genetic algorithm: NSGA-II," ', 'IEEE Trans. Evol. Comput.', ', vol. 6, no. 2, pp. 182-197, 2002.'),
    ('K. Deb and H. Jain, "An evolutionary many-objective optimization algorithm using reference-point-based nondominated sorting approach, part I: solving problems with box constraints," ', 'IEEE Trans. Evol. Comput.', ', vol. 18, no. 4, pp. 577-601, 2014.'),
    ('M. E. Baran and F. F. Wu, "Network reconfiguration in distribution systems for loss reduction and load balancing," ', 'IEEE Trans. Power Del.', ', vol. 4, no. 2, pp. 1401-1407, 1989.'),
    ('European Commission, Joint Research Centre, "Photovoltaic Geographical Information System (PVGIS)," online database. [Accessed: 2026].', '', ''),
]


def build():
    doc = Document()
    set_page(doc.sections[0]); set_cols(doc.sections[0], 1, 720)
    P(doc, TITLE, size=24, bold=False, align=AL.CENTER, before=0, after=3, line=0.96, md=False)
    P(doc, AUTHORS, size=11, align=AL.CENTER, after=0, md=False)
    for k, line in enumerate(AFFIL_LINES):
        P(doc, line, size=10, ital=True, align=AL.CENTER,
          after=(4 if k == len(AFFIL_LINES) - 1 else 0), md=False)

    new_section(doc, 2)
    ab = P(doc, "", after=3)
    font(ab.add_run("Abstract"), 9.0, bold=True, ital=True)
    font(ab.add_run("—"), 9.0, bold=True)
    runs_md(ab, ABSTRACT, 9.0, bold=True)
    it = P(doc, "", after=4)
    font(it.add_run("Index Terms"), 9.0, bold=True, ital=True)
    font(it.add_run("—"), 9.0, bold=True)
    font(it.add_run(INDEX_BODY), 9.0, bold=True)

    H1(doc, "I.  INTRODUCTION")
    for para in INTRO:
        P(doc, para, indent=True)

    H1(doc, "II.  PROPOSED METHOD")
    H2(doc, "A.  Differentiated Capacity-Load Ratio Model")
    P(doc, "Let j = 1, ..., K index the 110-kV substations of a county. The decision vector jointly sizes the "
           "capacity-load ratio and storage power rating of each substation,", indent=True)
    eq(doc, r"\mathbf{x}=\{(R_j,P_j)\}_{j=1}^{K}", "1")
    PT(doc, D_VARS)
    eq(doc, r"L_j^{\mathrm{p}}=L_j(1+g)^{n}", "2")
    PT(doc, AFTER_L)
    eq(doc, r"\tilde{r}_j(P_j)=\max\{0,\;r_j^{0}-P_j\}", "3")
    PT(doc, AFTER_R)
    eq(doc, r"1.8\le \sum_j R_j L_j^{\mathrm{p}}\,/\,\bigl(\delta\textstyle\sum_j L_j^{\mathrm{p}}\bigr)\le 2.2", "4")
    PT(doc, AFTER_AGG)
    eq(doc, [r"\tilde{r}_j(P_j)\le \beta\,R_j L_j^{\mathrm{p}}",
             r"\eta\sum_j \tilde{r}_j(P_j)\le \beta\,R_{220}\,\delta\sum_j L_j^{\mathrm{p}}"], "5")
    PT(doc, AFTER_UP)
    P(doc, OBJ_INTRO, indent=True)
    eq(doc, [r"\min_{\mathbf{x}}\ \bigl(f_1(\mathbf{x}),\;f_2(\mathbf{x})\bigr)",
             r"f_1(\mathbf{x})=\sum_j\bigl(c_R R_j L_j^{\mathrm{p}}+c_P P_j+C_j^{\text{loss}}+C_j^{\text{om}}\bigr)",
             r"f_2(\mathbf{x})=\sum_j \text{EENS}_j(R_j,T_j,L_j^{\mathrm{p}})"], "6")
    PT(doc, AFTER_F1)
    eq(doc, r"\text{EENS}_j=\sum_{s\in \mathcal{S}_j} p_s h_s\,\max\{0,\;L_j^{\mathrm{p}}-S_{j,s}\}", "7")
    PT(doc, AFTER_EENS)

    H2(doc, "B.  Engineering-Prior-Guided Search")
    P(doc, III_A, indent=True)
    algo_box(doc, "Algorithm 1: EPG-NSGA-II", ALGO1)
    PT(doc, III_B, indent=True)
    P(doc, III_C, indent=True)

    H1(doc, "III.  EXPERIMENTS")
    H2(doc, "A.  Setup")
    P(doc, SETUP, indent=True)
    H2(doc, "B.  Efficiency and Ablation")
    caption(doc, TAB_CAP, keep_next=True)
    three_line_table(doc, TAB_HEAD, TAB_ROWS, size=8.0, col_widths=[2.4, 1.45, 1.45, 1.45, 1.45])
    for para in EFF:
        P(doc, para, indent=True)

    H2(doc, "C.  County-Level Application at K=20")
    P(doc, APP, indent=True)

    fp = doc.add_paragraph(); fp.alignment = AL.CENTER; fp.paragraph_format.space_before = Pt(2)
    if os.path.exists(FIG):
        fp.add_run().add_picture(FIG, width=Cm(6.9))
    else:
        font(fp.add_run("[missing figure]"), 9, bold=True, color=RGBColor(0xAF, 0x3A, 0x3E))
    caption(doc, FIG_CAP, before=2)

    H1(doc, "IV.  CONCLUSION")
    P(doc, CONC, indent=True)

    H1(doc, "REFERENCES")
    for i, (pre, jour, post) in enumerate(REFS, 1):
        p = doc.add_paragraph(); p.alignment = AL.JUSTIFY
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 0.83
        p._p.get_or_add_pPr().get_or_add_ind().set(qn('w:hanging'), '150')
        font(p.add_run(f"[{i}]  "), 7.6)
        font(p.add_run(pre), 7.6)
        font(p.add_run(jour), 7.6, ital=True)
        font(p.add_run(post), 7.6)

    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    build()
