#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""论文骨架 Word —— 仿 IEEE 会议模版 conference-template-letter 的版面：
US Letter、L/R 0.62in·T 0.75in·B 1.0in、标题+作者跨栏、正文双栏、章标题小型大写居中、
子标题斜体、正文 10pt Times/宋体；宽图(框架/收敛)与 10 列 Table I 用整幅(跨双栏)放置。
逻辑线以每节 [逻辑线：…] 提示句形式体现在骨架里；正文只给核心逻辑线 + [待填] 占位。
文本经 de-ai-humanizer（论文·中度）：句长交错、去套路词、破折号「——」保留；术语/数据不动。作者 XH。
用法：python3 build_skeleton_docx.py"""
from __future__ import annotations
import os
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = "/home/roscy/ws_HengXU/徐州地区分布式新能源高渗透率地区110kV电网容载比弹性指标优化研究"
FIGD = os.path.join(ROOT, "实验/有EA/results/figures")
RESD = os.path.join(ROOT, "实验/有EA/results")
OUT = os.path.join(ROOT, "paper/徐州110kV容载比_MIND2026论文骨架_XH.docx")
SONG, TNR = "宋体", "Times New Roman"
GRAY = RGBColor(0x8A, 0x93, 0x9E)
COLW, FULLW = 8.4, 17.4   # 栏宽 / 整幅宽（cm）


def _set_font(run, size, bold=False, ital=False, ea=SONG, color=None, smallcaps=False):
    run.font.size = Pt(size); run.font.bold = bold; run.font.italic = ital
    run.font.name = TNR; run.font.small_caps = smallcaps
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    rf.set(qn('w:eastAsia'), ea); rf.set(qn('w:ascii'), TNR); rf.set(qn('w:hAnsi'), TNR)


def _runs_md(p, text, size, **kw):
    for i, seg in enumerate(text.split("**")):
        if seg:
            _set_font(p.add_run(seg), size, bold=(i % 2 == 1) or kw.get("bold", False),
                      ital=kw.get("ital", False), color=kw.get("color"))


def set_page(sec):
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.left_margin = sec.right_margin = Inches(0.62)
    sec.top_margin = Inches(0.75); sec.bottom_margin = Inches(1.0)


def set_cols(sec, num, space=360):
    sectPr = sec._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols'); sectPr.append(cols)
    cols.set(qn('w:num'), str(num)); cols.set(qn('w:space'), str(space)); cols.set(qn('w:equalWidth'), '1')


def new_section(doc, num):
    s = doc.add_section(WD_SECTION.CONTINUOUS); set_page(s); set_cols(s, num, 360 if num == 2 else 720)
    return s


def P(doc, text="", size=10, bold=False, ital=False, align=AL.JUSTIFY, before=0, after=3,
      line=1.06, indent=False, color=None, md=True):
    p = doc.add_paragraph(); p.alignment = align
    pf = p.paragraph_format; pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = line
    if indent:
        p._p.get_or_add_pPr().get_or_add_ind().set(qn('w:firstLineChars'), '200')
    if text:
        if md and not ital:
            _runs_md(p, text, size, bold=bold, color=color)
        else:
            _set_font(p.add_run(text), size, bold=bold, ital=ital, color=color)
    return p


def H1(doc, cn, en):
    p = doc.add_paragraph(); p.alignment = AL.CENTER
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    _set_font(p.add_run(f"{cn}  {en}"), 10, smallcaps=True)


def H2(doc, text):
    p = doc.add_paragraph(); p.alignment = AL.LEFT
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(1)
    _set_font(p.add_run(text), 10, ital=True)


def cue(doc, text):
    """逻辑线提示句（灰斜体，填充时删除）——把贯穿逻辑线落到每节形式上。"""
    P(doc, "［逻辑线：" + text + "］", size=8.5, ital=True, color=GRAY, after=2, md=False)


def stub(doc, text):
    P(doc, "［" + text + "］", size=8.5, ital=True, color=GRAY, after=3, md=False)


def _caption(doc, cn, en):
    P(doc, cn, size=8.5, bold=True, align=AL.CENTER, before=2, after=0, md=False)
    P(doc, en, size=8.5, bold=True, align=AL.CENTER, after=4, md=False)


def fig_col(doc, png, cn, en, w=COLW):
    p = doc.add_paragraph(); p.alignment = AL.CENTER; p.paragraph_format.space_before = Pt(4)
    if os.path.exists(png):
        p.add_run().add_picture(png, width=Cm(w))
    else:
        _set_font(p.add_run("［缺图：" + os.path.basename(png) + "］"), 9, bold=True, color=RGBColor(0xAF, 0x3A, 0x3E))
    _caption(doc, cn, en)


def fig_full(doc, png, cn, en, w=FULLW):
    new_section(doc, 1)
    fig_col(doc, png, cn, en, w=w)
    new_section(doc, 2)


# ---- 动态表 ----
def table1_rows():
    df = pd.read_csv(os.path.join(RESD, "ea_converge.csv"))
    mg = df.groupby("algo").gen.max().to_dict()
    order = ["classic", "nsga3", "repair_only", "warmstart_only", "fg"]
    rows = []
    for K in sorted(df.K.unique()):
        s = df[df.K == K]
        def hv(a):
            return s[(s.algo == a) & (s.gen == mg[a])].hv.mean()
        v = {a: hv(a) for a in order}
        fgc = s[s.algo == "fg"].groupby("gen").hv.mean()
        sp = {}
        for b in ("classic", "nsga3"):
            r = fgc[fgc >= hv(b)]; g = int(r.index.min()) if len(r) else mg["fg"]; sp[b] = mg[b] / g
        rows.append([str(int(K))] + [f"{v[a]:.3g}" for a in order]
                    + [f"{v['fg']/v['classic']:.3f}", f"{v['fg']/v['nsga3']:.3f}",
                       f"{sp['classic']:.1f}×", f"{sp['nsga3']:.1f}×"])
    return rows


def table2_rows():
    bl = pd.read_csv(os.path.join(RESD, "ea_baseline.csv")).iloc[0]
    rob = pd.read_csv(os.path.join(RESD, "ea_robustness.csv")).saving_pct
    band = pd.read_csv(os.path.join(RESD, "ea_band_effect.csv"))
    m, sd = rob.mean(), rob.std()
    bands = " / ".join(f"{r.band}:{r.min_cost_wan:.0f}" for _, r in band.iterrows())
    return [
        ["相对全县统一容载比 2.0（同风险）", f"{bl.rigid_f1_wan:.0f} → {bl.ea_f1_at_same_risk_wan:.0f} 万/年，省 {bl.ea_cost_saving_pct:.1f}%"],
        ["拐点方案（K=20）", "成本 18882 万/年、N-1 风险 114 MWh、聚合容载比 1.80（<2.0）"],
        ["一站一策路线", "11 站配储 / 9 站低容载比+强联络 / 0 站增容（最大 R*≈1.96）"],
        ["容载比带效应（成本下限, EENS=0）", bands + "（单调↑）"],
        [f"多县鲁棒（{len(rob)} 县重采样）", f"省 {m:.1f}%±{sd:.1f}%（{rob.min():.1f}–{rob.max():.1f}%，全为正）"],
        ["上级系统潮流约束", "加约束成本 Δ≈+25 万（<0.3%），近乎不绑定"],
    ]


def _borders(cell, top=None, bottom=None):
    tcPr = cell._tc.get_or_add_tcPr(); b = tcPr.find(qn('w:tcBorders'))
    if b is None:
        b = OxmlElement('w:tcBorders'); tcPr.append(b)
    for edge, sz in (('top', top), ('bottom', bottom)):
        if sz is None:
            continue
        el = b.find(qn('w:' + edge))
        if el is None:
            el = OxmlElement('w:' + edge); b.append(el)
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), str(sz)); el.set(qn('w:space'), '0'); el.set(qn('w:color'), '000000')


def add_table(doc, headers, rows, fs=8, first_left=False):
    t = doc.add_table(rows=1, cols=len(headers)); t.autofit = True
    for ci, h in enumerate(headers):
        t.rows[0].cells[ci].text = str(h)
    for rd in rows:
        cells = t.add_row().cells
        for ci, val in enumerate(rd):
            cells[ci].text = "" if val is None else str(val)
    nr = len(t.rows)
    for ci in range(len(headers)):
        _borders(t.rows[0].cells[ci], top=12, bottom=6); _borders(t.rows[nr - 1].cells[ci], bottom=12)
    for ri, row in enumerate(t.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = AL.CENTER; p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
                for r in p.runs:
                    _set_font(r, fs, bold=(ri == 0))
    if first_left:
        for ri in range(1, nr):
            for p in t.rows[ri].cells[0].paragraphs:
                p.alignment = AL.LEFT
    return t


def tbl_title(doc, cn, en):
    P(doc, cn, size=8.5, bold=True, align=AL.CENTER, before=6, after=0, md=False)
    P(doc, en, size=8.5, bold=True, align=AL.CENTER, after=1, md=False)


# ==================== 构建 ====================
doc = Document()
set_page(doc.sections[0]); set_cols(doc.sections[0], 1)   # 标题块：跨栏单栏

P(doc, "分布式新能源高渗透地区县级 110 kV 电网容载比弹性的多目标进化优化",
  size=20, bold=True, align=AL.CENTER, after=2, md=False)
P(doc, "Multi-Objective Evolutionary Optimization of Capacity-Load-Ratio Elasticity "
       "for County-Level 110 kV Grids with High-Penetration Distributed Renewables",
  size=11, bold=True, align=AL.CENTER, after=3, md=False)
stub(doc, "作者、单位、邮箱待填（IEEE 作者块：Given Name Surname / dept. / org. / city, country / email）")

new_section(doc, 2)   # ↓ 正文双栏

P(doc, "", after=0)
p = doc.paragraphs[-1]; _set_font(p.add_run("摘要 Abstract—"), 9, bold=True)
_runs_md(p, "分布式光伏高渗透使县域主变在午间净反送下逼近反向重载；现行全县统一容载比 2.0 要么过度投资，"
            "要么留下隐患。本文提出容载比“向下弹性”，并把它形式化为县级多站联合的成本↔N-1 风险多目标优化——"
            "决策为每站容载比与储能，约束含逐站反送闸、县聚合容载比带与上级系统潮流。可行域随站数骤缩、经典 "
            "NSGA-II 前期收敛慢；为此提出可行性引导的 **FG-NSGA-II**，以可行性修复算子为主驱动、启发式热启动为早期加速。"
            "在公开数据合成的县级基准（K=20）上，方案相对统一容载比 2.0 在同风险下节省 **13.1%** 并给出一站一策；"
            "FG-NSGA-II 用约一半预算即达两条基线跑满预算的收敛质量，等价 **4.8–11.8×** 更少代数"
            "（充分收敛后基线大体追平，如实交代）。", 9)
p.alignment = AL.JUSTIFY
kp = doc.add_paragraph(); _set_font(kp.add_run("关键词 Keywords—"), 9, bold=True, ital=True)
_set_font(kp.add_run("容载比弹性；分布式光伏高渗透；多目标进化优化；NSGA-II；可行性修复算子；一站一策。"), 9, ital=True)

# I. 引言
H1(doc, "I.", "引言 Introduction")
cue(doc, "抛出问题与缺口——引出“把弹性建模成县级约束多目标优化 + 用 FG-NSGA-II 求解”这条主线的起点")
P(doc, "分布式光伏高渗透正在改变配电网的潮流方向：午间净反送把 110 kV 主变推向反向重载，"
       "容载比这一规划核心指标随之面临新的取舍。", indent=True)
stub(doc, "文献综述待补（编年体、密集引用 [n]）：一线为反送承载力与容载比导则，另一线为进化多目标优化在电网规划的应用")
P(doc, "现有做法多对全县取同一容载比，未利用站间异质与同时率；已有进化多目标优化的电网规划应用多为小规模、弱约束；"
       "经典算法的约束处理依赖惩罚与随机初始化，在强约束、可行域稀疏、随站数膨胀的问题上前期收敛慢、可行解稀。", indent=True)
P(doc, "本文把“容载比向下弹性”形式化为县级多站联合的约束多目标优化，并针对其问题结构提出 FG-NSGA-II。主要贡献：", indent=True)
P(doc, "**1)** 将容载比弹性建模为县级多站“成本↔N-1 风险”约束多目标优化，并诊断出“带与系统约束耦合、可行域随 K 稀疏化”正是通用 MOEA 在此失效之因；")
P(doc, "**2)** 针对该结构提出 FG-NSGA-II——以可行性修复为主驱动、启发式热启动为早期加速，使县级大规模寻优可解；")
P(doc, "**3)** 在公开数据合成的县级基准上给出双层证据：应用层相对统一容载比 2.0 省 13.1% 且一站一策，算法层以约一半预算胜过两条基线、等价 4.8–11.8× 更少代数。")
P(doc, "本文组织如下：第 II 节给出问题形式化与算例，第 III 节提出 FG-NSGA-II 及其伪代码，第 IV 节报告实验，第 V 节总结。", indent=True)

# II. 问题形式化
H1(doc, "II.", "问题形式化 Problem Formulation")
cue(doc, "把“弹性”落成可计算的约束多目标优化，并暴露“带+系统约束→不可分、强约束、随 K 膨胀”——为 III 的算法必要性铺垫")
H2(doc, "A. 决策变量与容载比弹性")
P(doc, "决策向量 x=[R₁..R_K, P₁..P_K]：每站容载比 R_j 与储能功率 P_j（储能容量 E=2h×P，从属）；“向下弹性”指在合规带内能降尽降。", indent=True)
stub(doc, "式 (1) 决策空间待填")
H2(doc, "B. 目标函数")
P(doc, "min f₁ = 县总年化成本（资本年金 + 双向网损 + 储能 + 运维 − 残值）；min f₂ = 县 N-1 缺供 ΣEENS（按未来峰荷）。"
       "并说明为何取“成本↔风险”而非“成本↔弃光”（后者在可行集内退化）。", indent=True)
stub(doc, "式 (2)(3) 目标待填；随后以 “the symbols have the following meaning:” 破折号符号表定义")
H2(doc, "C. 约束")
P(doc, "① 逐站过反向承载力硬闸；② 县聚合容载比带 [1.8,2.2]（同时率 0.85，核心耦合、使问题不可分）；"
       "③ 上级系统潮流（聚合反送 ≤ 上级反向限额，R₂₂₀=1.8）。", indent=True)
stub(doc, "式 (4)–(6) 约束待填")
H2(doc, "D. 合成基准算例")
P(doc, "以缩放 IEEE33 馈线为载体、PVGIS 三地辐照锚定，按四元组抽样合成县；主算例 K=20。声明：公开数据合成、非实测。", indent=True)

# III. 所提方法
H1(doc, "III.", "所提方法：FG-NSGA-II  Proposed Method")
cue(doc, "回应 II 暴露的结构——可行性由构造保证（修复为主、热启动为辅），这既是应用可解之钥，也是算法效率之源")
P(doc, "FG-NSGA-II 把确定性物理评价器与 NSGA-II 搜索器耦合，并把约束当作“处理”而非“求解”：物理模型只判定越限，"
       "修复算子把不可行解就近构造回可行域，进化只施加可行性优先的选择压力。", indent=True)
fig_full(doc, os.path.join(FIGD, "fig_framework.png"),
         "图 1  FG-NSGA-II 融合框架与可行性修复流程",
         "Fig. 1.  Overview of the FG-NSGA-II framework and the feasibility-repair procedure.")
H2(doc, "A. 融合框架")
P(doc, "物理评价器（双向 Z4-LCC，约 4 ms/次，无需代理模型）+ NSGA-II 搜索器；判定 / 修复 / 选择三分。", indent=True)
H2(doc, "B. 可行性修复算子（主驱动）")
P(doc, "一次前向、闭式地把不可行解投影到近似可行：缩放 R 入带 → 逐站补储或必要时升 R 过反送闸 → 补储满足系统约束。"
       "修复不调用目标函数，因而不占评估预算。", indent=True)
H2(doc, "C. 启发式热启动（早期加速器）")
P(doc, "以拉丁超立方保多样性、注入“低 R + 零储”种子并全部经修复，令初始种群即近似可行。", indent=True)
H2(doc, "D. 为何有效，又为何仍需进化")
P(doc, "大 K 下可行域骤缩，随机初始化难以命中；修复直接构造可行解绕过瓶颈，使 FG 用一半预算即追平基线满预算。"
       "但修复只给可行、给不出前沿，成本↔风险的权衡仍靠进化搜索。", indent=True)
stub(doc, "Algorithm 1（FG-NSGA-II 主循环伪代码）待填")

# IV. 实验研究
H1(doc, "IV.", "实验研究 Experimental Study")
cue(doc, "两层证据：B 证方案好用（省 13.1%、一站一策），C 证使能方法为何快且随规模仍快——共同回扣“结构即失效之因、亦即 FG 之用”")
H2(doc, "A. 实验设置")
P(doc, "规模 K∈{10,20,30,40}；基线为经典 NSGA-II 与 NSGA-III（种群严格相等），另含 repair-only / warm-start-only 两消融变体；"
       "指标以 HV 为主、小 K 用 IGD 对 Minkowski 真前沿；逐代记 HV（基线 gen800、FG 类 gen400），3 seed 取 mean±std，每 K 共同参考点。"
       "种群 pop=100 固定，故代数与目标函数评价次数（NFE=pop×gen）及墙钟时间成正比。", indent=True)
H2(doc, "B. 县级应用结果（K=20）")
P(doc, "拐点方案相对统一容载比 2.0 在同风险下省 13.1%，并给出一站一策：高发电荷比站配储为主、低者取低容载比配强联络，无站需 R>2.0。", indent=True)
fig_col(doc, os.path.join(FIGD, "fig_county_pareto.png"),
        "图 2(a)  县级成本–N-1 风险 Pareto 前沿与拐点",
        "Fig. 2(a).  Cost–N-1 risk Pareto front and the knee point.")
fig_col(doc, os.path.join(FIGD, "fig_strategy.png"),
        "图 2(b)  拐点处逐站“一站一策”",
        "Fig. 2(b).  Per-substation strategy at the knee point.")
tbl_title(doc, "表 II  应用结果汇总（FG-NSGA-II，K=20）", "TABLE II.  SUMMARY OF APPLICATION RESULTS (K=20)")
add_table(doc, ["项目", "结果"], table2_rows(), fs=8, first_left=True)
H2(doc, "C. 算法效率与组件消融（核心）")
P(doc, "FG-NSGA-II 以约一半预算即达两条基线跑满预算的收敛质量，等价 4.8–11.8× 更少代数（预算无关）。"
       "收敛质量差距小且如实披露：收敛 HV 比在 K≤30 时不足 1%、K=40 才 4.8%——本文主张的是效率，而非收敛质量更高。", indent=True)
fig_full(doc, os.path.join(FIGD, "fig_enhanced_convergence.png"),
         "图 3（核心）  各规模逐代 HV 收敛（4 个 K、5 变体 mean±std）",
         "Fig. 3.  Per-generation hypervolume (mean±std, 3 seeds) for K in {10,20,30,40}.")
fig_col(doc, os.path.join(FIGD, "fig_scaling_gap.png"),
        "图 4  iso-quality 加速比随规模 K",
        "Fig. 4.  Iso-quality speedup versus scale K.")
P(doc, "两组件消融：repair-only 的收敛 HV 与完整 FG 几乎重合（比值 0.997–1.000），说明修复算子是主驱动；"
       "warm-start-only 在早期超过经典（gen100 至 +6.3%）但收敛后被洗掉，说明热启动只提前收敛、不改终值。", indent=True)
# Table I 宽（10 列）→ 整幅
new_section(doc, 1)
tbl_title(doc, "表 I  5 变体收敛 HV、加速比与消融归因（K∈{10,20,30,40}）",
          "TABLE I.  CONVERGED HYPERVOLUME, ISO-QUALITY SPEEDUP AND ABLATION")
add_table(doc, ["K", "Classic", "NSGA-III", "Repair-only", "Warm-start", "FG (ours)",
                "FG/Cls", "FG/N3", "加速×Cls", "加速×N3"], table1_rows(), fs=8)
P(doc, "注：离散性以 3 seed 的 mean±std 报告；正式显著性检验（Wilcoxon）列为扩展版工作。", size=8, ital=True, color=GRAY, md=False)
new_section(doc, 2)

# V. 结论
H1(doc, "V.", "结论 Conclusion")
cue(doc, "收束主线：弹性建模 → FG-NSGA-II → 双层证据；再指向实测标定等未来工作")
P(doc, "本文把容载比“向下弹性”建模为县级多站约束多目标优化，提出 FG-NSGA-II 使其可解，并给出双层证据："
       "应用层相对统一容载比 2.0 省 13.1% 且一站一策，算法层以约一半预算、4.8–11.8× 更少代数达同等收敛质量。", indent=True)
stub(doc, "future work 待补：实测标定、储能能量约束、区域×规模真·跨县、上级约束更强绑定情形、更大规模与统计检验")

H1(doc, "", "参考文献 References")
stub(doc, "约 24 条 IEEE 格式待补：容载比/反送承载力导则、分布式光伏并网、NSGA-II 原文、NSGA-III、"
          "约束处理/修复算子 EMO、HV/IGD 指标、进化多目标在电网规划应用、pymoo；可复用无EA版 [14]–[30]")

cp = doc.core_properties
cp.author = "XH"; cp.last_modified_by = "XH"
cp.title = "县级110kV容载比弹性 FG-NSGA-II —— 论文骨架（IEEE 会议模版格式）"
doc.save(OUT)
print(f"[OK] {OUT}  ({os.path.getsize(OUT)/1024:.0f} KB)  author={cp.author}")
