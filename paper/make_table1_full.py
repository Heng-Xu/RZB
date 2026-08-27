#!/usr/bin/env python3
"""Standalone Table I (full comparison, NSGA-II = 1.000 baseline) + corrected
description paragraph, in the paper's IEEE three-line style, for pasting into
the hand-edited Word draft. Reuses helpers/data from build_paper_docx.py so it
stays in sync. Output: paper/table1_full.docx
"""
import os
from docx import Document
from docx.shared import Pt
import build_paper_docx as B

OUT = os.path.join(B.ROOT, "paper/table1_full.docx")


def main():
    doc = Document()
    B.set_page(doc.sections[0])
    B.caption(doc, B.TAB_CAP, keep_next=True)
    B.three_line_table(doc, B.TAB_HEAD, B.TAB_ROWS, size=9.0,
                       col_widths=[2.6, 1.6, 1.6, 1.6, 1.6])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    B.font(p.add_run(B.EFF[0]), 9.5)

    # Reference to add for the NSGA-III citation (pull from the paper's REFS so it stays in sync)
    pre, jour, post = next(r for r in B.REFS if "Jain" in r[0])
    r = doc.add_paragraph()
    r.paragraph_format.space_before = Pt(10)
    B.font(r.add_run("Reference to add (cite for NSGA-III):"), 9.0, bold=True)
    r2 = doc.add_paragraph()
    r2.paragraph_format.space_before = Pt(2)
    B.font(r2.add_run(pre), 9.0)
    B.font(r2.add_run(jour), 9.0, ital=True)
    B.font(r2.add_run(post), 9.0)
    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    main()
